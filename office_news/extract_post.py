"""
네이버 블로그 단일 포스트 추출기
  - 텍스트 → Word (.docx)
  - 이미지 → 개별 파일 저장
"""
import os, re, time
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from urllib.parse import urlparse

BLOG_ID    = "wons0_0"
LOG_NO     = "224168110325"
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output", f"{BLOG_ID}_{LOG_NO}")
POST_URL   = f"https://blog.naver.com/{BLOG_ID}/{LOG_NO}"

SESSION = requests.Session()
SESSION.headers.update({
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Referer": POST_URL,
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8",
})

SKIP_DOMAINS = ("staticmap", ".bin", "tracking", "icon", "btn_", "blogimgs",
                "ssl.pstatic", "blogpfthumb", "spc.gif", "nblog")


def fetch(url: str) -> str:
    r = SESSION.get(url, timeout=20)
    r.raise_for_status()
    r.encoding = "utf-8"
    return r.text


def clean_img_url(url: str) -> str:
    return url if "postfiles.pstatic.net" in url else url.split("?")[0]


def ext_from_url(url: str) -> str:
    ext = os.path.splitext(urlparse(url).path)[1].lower()
    return ext if ext in (".jpg", ".jpeg", ".png", ".gif", ".webp") else ".jpg"


def is_valid_img(url: str, tag) -> bool:
    lower = url.lower()
    if any(x in lower for x in SKIP_DOMAINS):
        return False
    if not url.startswith("http"):
        return False
    try:
        if tag.get("width")  and int(tag["width"])  < 50: return False
        if tag.get("height") and int(tag["height"]) < 50: return False
    except (ValueError, TypeError):
        pass
    return True


def download_img(url: str, dest: str) -> bool:
    try:
        r = SESSION.get(url, timeout=20, stream=True,
                        headers={"Accept": "image/avif,image/webp,image/apng,image/*,*/*"})
        r.raise_for_status()
        with open(dest, "wb") as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
        return True
    except Exception as e:
        print(f"  [skip] {url[:70]} — {e}")
        return False


def get_title(soup: BeautifulSoup) -> str:
    for tag, attrs in [
        ("h3",  {"class": "se-title-text"}),
        ("div", {"class": "se-title-text"}),
        ("div", {"class": "post-title"}),
        ("h1",  {"class": "htitle"}),
    ]:
        t = soup.find(tag, attrs)
        if t:
            return t.get_text(strip=True)
    t = soup.find("title")
    if t:
        return re.sub(r"\s*[:|-]?\s*네이버\s*블로그.*$", "", t.get_text(strip=True)).strip()
    return f"{BLOG_ID}_{LOG_NO}"


def safe_name(text: str) -> str:
    text = re.sub(r'[\\/:*?"<>|]', "", text)
    return re.sub(r"\s+", "_", text.strip())[:60]


def collect_content(body) -> list:
    items, seen_imgs, seen_texts, counter = [], set(), set(), 0
    for elem in body.descendants:
        if not hasattr(elem, "name") or elem.name is None:
            continue
        if elem.name == "img":
            src = elem.get("data-lazy-src") or elem.get("src") or ""
            if not src or src.startswith("data:"):
                continue
            src = clean_img_url(src)
            if src in seen_imgs or not is_valid_img(src, elem):
                continue
            seen_imgs.add(src)
            counter += 1
            items.append(("image", src, counter))
        elif elem.name in ("p","h1","h2","h3","h4","h5","h6","li","span","div"):
            if elem.find(["p","h1","h2","h3","div","li"]) or elem.find("img"):
                continue
            text = elem.get_text(separator=" ", strip=True)
            if not text or len(text) <= 1:
                continue
            key = re.sub(r"\s+", " ", text).strip()
            if key in seen_texts:
                continue
            seen_texts.add(key)
            items.append(("text", text))
    return items


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"포스트 로딩 중… {POST_URL}")

    iframe_url = (
        f"https://blog.naver.com/PostView.naver"
        f"?blogId={BLOG_ID}&logNo={LOG_NO}&redirect=Dlog&widgetTypeCall=true"
    )
    html = fetch(iframe_url)
    if len(html) < 500:
        print("  iframe 로딩 실패, 모바일 URL 시도...")
        html = fetch(f"https://m.blog.naver.com/{BLOG_ID}/{LOG_NO}")

    print(f"HTML 크기: {len(html):,} bytes")

    soup  = BeautifulSoup(html, "lxml")
    title = get_title(soup)
    base  = safe_name(title)
    print(f"제목: {title}")

    body = (
        soup.find("div", class_="se-main-container")
        or soup.find("div", id="postViewArea")
        or soup.find("div", class_="post-view")
        or soup.find("div", id="content")
        or soup.body
    )
    content   = collect_content(body or soup)
    img_items = [(item[1], item[2]) for item in content if item[0] == "image"]
    print(f"발견된 이미지: {len(img_items)}개")

    # 이미지 다운로드
    img_files: dict[int, str] = {}
    for url, idx in img_items:
        fname = f"{base}_이미지_{idx:02d}{ext_from_url(url)}"
        dest  = os.path.join(OUTPUT_DIR, fname)
        print(f"  [{idx}/{len(img_items)}] {url[:80]}")
        if download_img(url, dest):
            img_files[idx] = fname
        time.sleep(0.3)

    # Word 문서
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"원본 URL: {POST_URL}")
    doc.add_paragraph("")
    text_count = 0
    for item in content:
        if item[0] == "image":
            fname = img_files.get(item[2], f"이미지_{item[2]:02d}")
            p = doc.add_paragraph()
            p.add_run(f"[{os.path.splitext(fname)[0]}]").bold = True
        else:
            p = doc.add_paragraph(item[1])
            p.style.font.size = Pt(11)
            text_count += 1

    docx_name = f"{base}_텍스트.docx"
    doc.save(os.path.join(OUTPUT_DIR, docx_name))

    print(f"\n=== 완료 ===")
    print(f"저장 폴더  : {OUTPUT_DIR}")
    print(f"이미지     : {len(img_files)}개")
    print(f"텍스트단락 : {text_count}개")
    print(f"Word 문서  : {docx_name}")


if __name__ == "__main__":
    main()
