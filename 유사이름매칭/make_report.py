from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 페이지 여백 설정
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is not None:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), val.get('sz', '4'))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))  # 1cm = 567 twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

# ── 로고 & 제목 ──────────────────────────────────────
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_run = logo_p.add_run('JU  정원유니어스')
logo_run.font.size = Pt(11)
logo_run.font.bold = True
logo_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xB0)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('교육 결과 보고서')
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
title_p.paragraph_format.space_before = Pt(4)
title_p.paragraph_format.space_after = Pt(6)

# ── 결재 테이블 ──────────────────────────────────────
# 열: 결재 | 담당 | 팀장 | 본부장 | 대표이사
approval = doc.add_table(rows=2, cols=5)
approval.alignment = WD_TABLE_ALIGNMENT.CENTER
approval.style = 'Table Grid'

# 열 너비 설정
widths = [Cm(1.5), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8)]
for row in approval.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

# 헤더 행
headers = ['결  재', '담  당', '팀  장', '본부장', '대표이사']
for i, h in enumerate(headers):
    cell = approval.cell(0, i)
    cell.text = h
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(9)
    run.font.bold = True
    set_cell_bg(cell, 'D9E1F2')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# 이름 행
names = ['', '박민석', '강호진', '', '임재헌']
for i, name in enumerate(names):
    cell = approval.cell(1, i)
    cell.text = name
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# "결재" 셀 세로 병합 (행 0 col 0 → 행 1 col 0)
approval.cell(0, 0).merge(approval.cell(1, 0))
approval.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run0 = approval.cell(0, 0).paragraphs[0].runs[0]
run0.font.size = Pt(9)
run0.font.bold = True
approval.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

set_row_height(approval.rows[0], 0.7)
set_row_height(approval.rows[1], 1.2)

doc.add_paragraph()  # 간격

# ── 섹션 헤더 스타일 함수 ──────────────────────────────
def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    # 하단 테두리
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A56B0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_info_table(doc, rows_data):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    col_widths = [Cm(2.5), Cm(12.0)]
    for row_idx, (label, value) in enumerate(rows_data):
        row = tbl.rows[row_idx]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        # 레이블
        c0 = row.cells[0]
        c0.text = label
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c0.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_bg(c0, 'EEF2FA')
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 값
        c1 = row.cells[1]
        c1.text = value
        run1 = c1.paragraphs[0].runs[0]
        run1.font.size = Pt(10)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_row_height(row, 0.7)
    doc.add_paragraph()

def add_bullet(doc, text, sub=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if sub:
        run2 = p.add_run(f'\n{sub}')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# ── 1. 교육 개요 ──────────────────────────────────────
add_section_title(doc, '1. 교육 개요')
add_info_table(doc, [
    ('교육명', 'AI 기반 스마트 오피스 활용 실무 과정'),
    ('일 시', '2026년 4월 20일 ~ 22일, 3일간'),
    ('장 소', '한국생산성본부'),
    ('참석자', '마케팅팀 이준혁 대리'),
])

# ── 2. 교육 내용 요약 ─────────────────────────────────
add_section_title(doc, '2. 교육 내용 요약')
add_bullet(doc, 'AI 도구(ChatGPT, Midjourney) 활용 실무',
           '(트렌드 분석, 보고서 작성, 이미지 생성 실습)')
add_bullet(doc, '데이터 분석 및 시각화',
           '(Excel AI 기능, 효과적인 시각화 기법)')
add_bullet(doc, '자동화 프로세스',
           '(RPA 기초, 단순 반복 업무 자동화)')
doc.add_paragraph()

# ── 3. 업무 적용 부분 ─────────────────────────────────
add_section_title(doc, '3. 업무 적용 부분')
add_bullet(doc, '제안서 및 보고서 작성 속도 향상', '(AI 기반 초안 작성)')
add_bullet(doc, '시장 데이터 분석 강화', '(효율적인 트렌드 파악)')
add_bullet(doc, '고객 응대 자동화 (이메일, SNS)', '(생산성 증대)')
doc.add_paragraph()

# ── 4. 향후 계획 ──────────────────────────────────────
add_section_title(doc, '4. 향후 계획')
plan_p = doc.add_paragraph()
plan_run = plan_p.add_run(
    '교육 이수 후 팀 내 공유 세션을 진행하고, 습득한 AI 도구를 실무 프로젝트에 단계적으로 적용할 계획입니다.'
)
plan_run.font.size = Pt(10)
plan_p.paragraph_format.space_after = Pt(16)

# ── 제출 정보 ──────────────────────────────────────────
submit_p = doc.add_paragraph()
submit_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
submit_run = submit_p.add_run('2026년 4월 27일\n제출자: 이준혁  (인)')
submit_run.font.size = Pt(10)

# 저장
out_path = r'C:\Users\JW\Desktop\교육결과보고서.docx'
doc.save(out_path)
print(f'저장 완료: {out_path}')
