"""
Naver blog extractor — saves images as files and text as a Word document.
"""
import os
import re
import time
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from urllib.parse import urljoin, urlparse, unquote

BLOG_ID = "yesclinic_juan"
LOG_NO = "223143230135"
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "naver_output")
POST_URL = f"https://blog.naver.com/{BLOG_ID}/{LOG_NO}"
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Referer": POST_URL,
    "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8",
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8",
}
SESSION = requests.Session()
SESSION.headers.update(HEADERS)


def fetch_html(url: str) -> str:
    resp = SESSION.get(url, timeout=15)
    resp.raise_for_status()
    resp.encoding = "utf-8"
    return resp.text


def get_post_html() -> str:
    """Naver blog content lives inside an iframe — fetch it directly."""
    iframe_url = (
        f"https://blog.naver.com/PostView.naver"
        f"?blogId={BLOG_ID}&logNo={LOG_NO}&redirect=Dlog&widgetTypeCall=true"
    )
    html = fetch_html(iframe_url)
    if len(html) < 500:
        # Fallback: try mobile
        html = fetch_html(f"https://m.blog.naver.com/{BLOG_ID}/{LOG_NO}")
    return html


def clean_img_url(url: str) -> str:
    """Keep Naver resize params for postfiles; drop query string for others."""
    if "postfiles.pstatic.net" in url:
        return url  # keep ?type=w966 — required for access
    return url.split("?")[0]


def download_image(url: str, dest_path: str) -> bool:
    try:
        r = SESSION.get(url, timeout=15, stream=True)
        r.raise_for_status()
        with open(dest_path, "wb") as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
        return True
    except Exception as e:
        print(f"  [skip] {url[:70]} - {e}")
        return False


def ext_from_url(url: str) -> str:
    path = urlparse(url).path
    ext = os.path.splitext(path)[1].lower()
    if ext in (".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"):
        return ext
    return ".jpg"


def is_valid_image_url(url: str) -> bool:
    """Skip map tiles, tracking pixels, and non-image binaries."""
    lower = url.lower()
    if any(x in lower for x in ("staticmap", ".bin", "tracking", "icon", "btn_")):
        return False
    return True


def safe_filename(title: str) -> str:
    """Convert a title to a safe filename (strip illegal chars, limit length)."""
    title = re.sub(r'[\\/:*?"<>|]', "", title)  # strip Windows-illegal chars
    title = re.sub(r"\s+", "_", title.strip())
    return title[:60]  # cap at 60 chars


def get_title(soup: BeautifulSoup) -> str:
    """Extract the post title from the parsed HTML."""
    # Smart Editor 3
    tag = soup.find("h3", class_="se-title-text") or soup.find("div", class_="se-title-text")
    if tag:
        return tag.get_text(strip=True)
    # Older editor
    tag = soup.find("div", class_="post-title") or soup.find("h1", class_="htitle")
    if tag:
        return tag.get_text(strip=True)
    # <title> tag as fallback
    tag = soup.find("title")
    if tag:
        text = tag.get_text(strip=True)
        # Strip " : 네이버 블로그" suffix if present
        text = re.sub(r"\s*[:|-]?\s*네이버\s*블로그.*$", "", text).strip()
        return text
    return f"{BLOG_ID}_{LOG_NO}"


def collect_content(body) -> list:
    """Walk the body DOM in document order and return an ordered list of
    ('image', url, img_index) and ('text', text) items."""
    items = []
    seen_imgs = set()
    seen_texts = set()
    img_counter = 0

    for elem in body.descendants:
        if not hasattr(elem, "name") or elem.name is None:
            continue

        if elem.name == "img":
            src = elem.get("data-lazy-src") or elem.get("src") or ""
            if not src or src.startswith("data:"):
                continue
            src = clean_img_url(src)
            if src in seen_imgs or not is_valid_image_url(src):
                continue
            w, h = elem.get("width", ""), elem.get("height", "")
            try:
                if w and int(w) < 50:
                    continue
                if h and int(h) < 50:
                    continue
            except ValueError:
                pass
            seen_imgs.add(src)
            img_counter += 1
            items.append(("image", src, img_counter))

        elif elem.name in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "span", "div"):
            # Skip containers that have block children (avoid duplicating text)
            if elem.find(["p", "h1", "h2", "h3", "div", "li"]):
                continue
            # Skip elements that hold an image (placeholder handles those)
            if elem.find("img"):
                continue
            text = elem.get_text(separator=" ", strip=True)
            if not text or len(text) <= 1:
                continue
            key = re.sub(r"\s+", " ", text).strip()
            if key in seen_texts:
                continue
            seen_texts.add(key)
            items.append(("text", text))

    return items


def extract(html: str):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    soup = BeautifulSoup(html, "lxml")

    # ── extract title for filenames ────────────────────────────────────────────
    title = get_title(soup)
    base_name = safe_filename(title)
    print(f"블로그 제목: {title}")

    # ── locate main post body ──────────────────────────────────────────────────
    body = (
        soup.find("div", class_="se-main-container")
        or soup.find("div", id="postViewArea")
        or soup.find("div", class_="post-view")
        or soup.find("div", id="content")
        or soup.body
    )

    # ── collect all content in document order ──────────────────────────────────
    content = collect_content(body or soup)

    img_items = [(item[1], item[2]) for item in content if item[0] == "image"]
    print(f"발견된 이미지: {len(img_items)}개")

    # ── download images ────────────────────────────────────────────────────────
    img_files: dict[int, str] = {}
    for url, idx in img_items:
        ext = ext_from_url(url)
        fname = f"{base_name}_이미지_{idx:02d}{ext}"
        dest = os.path.join(OUTPUT_DIR, fname)
        print(f"  [{idx}/{len(img_items)}] 다운로드 중... {url[:80]}")
        if download_image(url, dest):
            img_files[idx] = fname
        time.sleep(0.3)

    # ── build Word document in document order ──────────────────────────────────
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"원본 URL: {POST_URL}")
    doc.add_paragraph("")

    text_count = 0
    for item in content:
        if item[0] == "image":
            _, url, idx = item
            fname = img_files.get(idx, f"이미지_{idx:02d}")
            label = os.path.splitext(fname)[0]
            p = doc.add_paragraph()
            run = p.add_run(f"[{label}]")
            run.bold = True
        else:
            _, text = item
            p = doc.add_paragraph(text)
            p.style.font.size = Pt(11)
            text_count += 1

    docx_name = f"{base_name}_텍스트.docx"
    doc.save(os.path.join(OUTPUT_DIR, docx_name))

    # ── summary ────────────────────────────────────────────────────────────────
    print(f"\n=== 완료 ===")
    print(f"저장 폴더 : {OUTPUT_DIR}")
    print(f"이미지    : {len(img_files)}개 저장됨")
    print(f"텍스트    : {text_count}개 단락")
    print(f"Word 문서 : {docx_name}")
    if img_files:
        print("\n저장된 이미지 파일:")
        for fname in sorted(img_files.values()):
            print(f"  {fname}")


if __name__ == "__main__":
    print("페이지 로딩 중…")
    html = get_post_html()
    print(f"HTML 크기: {len(html):,} bytes")
    extract(html)
