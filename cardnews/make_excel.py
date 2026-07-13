"""
naver_output_2023 폴더를 읽어 날짜 / 제목 / 요약 엑셀 파일을 생성합니다.
요약은 각 Word 문서의 도입부 본문 텍스트를 합쳐 만듭니다.
"""
import os
import re
import sys
from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

INPUT_ROOT = os.path.join(os.path.dirname(__file__), "naver_output_2023")
OUTPUT_XLS = os.path.join(os.path.dirname(__file__), "naver_2023_목록.xlsx")

SKIP_PREFIXES = ("원본 URL:", "작성일:", "http")
SKIP_BRACKET  = re.compile(r"^\[.+\]$")   # image placeholders like [제목_이미지_01]


def extract_summary(docx_path: str, max_chars: int = 120) -> str:
    """Word 문서에서 이미지 플레이스홀더와 헤더를 제외한 본문 앞부분을 반환."""
    try:
        doc = Document(docx_path)

        # 첫 번째 헤딩에서 제목 텍스트 추출 (중복 제거용)
        heading_text = ""
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                heading_text = para.text.strip()
                break

        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                continue
            if any(text.startswith(p) for p in SKIP_PREFIXES):
                continue
            if SKIP_BRACKET.match(text):
                continue
            # 제목과 동일한 텍스트 단락 제외 (블로그 본문 상단 제목 반복)
            if heading_text and text == heading_text:
                continue
            parts.append(text)
            if sum(len(p) for p in parts) >= max_chars:
                break

        summary = " ".join(parts)
        if len(summary) > max_chars:
            summary = summary[:max_chars].rsplit(" ", 1)[0] + "…"
        return summary
    except Exception as e:
        return f"(읽기 오류: {e})"


def parse_folder(folder_name: str):
    """'2023-12-29_제목' 형식에서 날짜와 제목 분리."""
    m = re.match(r"^(\d{4}-\d{2}-\d{2})_(.*)", folder_name)
    if m:
        return m.group(1), m.group(2).replace("_", " ")
    return "", folder_name.replace("_", " ")


def main():
    folders = sorted(
        f for f in os.listdir(INPUT_ROOT)
        if os.path.isdir(os.path.join(INPUT_ROOT, f))
    )
    print(f"폴더 수: {len(folders)}개")

    rows = []
    for folder in folders:
        fp = os.path.join(INPUT_ROOT, folder)
        date, title = parse_folder(folder)

        # find the .docx file
        docx_files = [f for f in os.listdir(fp) if f.endswith(".docx")]
        if not docx_files:
            summary = "(Word 파일 없음)"
        else:
            summary = extract_summary(os.path.join(fp, docx_files[0]))

        rows.append((date, title, summary))
        print(f"  {date}  {title[:30]}")

    # ── build Excel ────────────────────────────────────────────────────────────
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "2023년 블로그 목록"

    # header style
    header_font    = Font(name="맑은 고딕", bold=True, size=11, color="FFFFFF")
    header_fill    = PatternFill("solid", fgColor="2E75B6")
    header_align   = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_align     = Alignment(vertical="center", wrap_text=True)
    thin           = Side(style="thin", color="BFBFBF")
    border         = Border(left=thin, right=thin, top=thin, bottom=thin)

    headers = ["날짜", "제목", "요약"]
    col_widths = [13, 40, 70]

    ws.row_dimensions[1].height = 28
    for col, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font   = header_font
        cell.fill   = header_fill
        cell.alignment = header_align
        cell.border = border
        ws.column_dimensions[get_column_letter(col)].width = w

    # data rows
    for r, (date, title, summary) in enumerate(rows, 2):
        ws.row_dimensions[r].height = 40
        for col, val in enumerate([date, title, summary], 1):
            cell = ws.cell(row=r, column=col, value=val)
            cell.font      = Font(name="맑은 고딕", size=10)
            cell.alignment = cell_align
            cell.border    = border
        # alternate row shading
        if r % 2 == 0:
            for col in range(1, 4):
                ws.cell(row=r, column=col).fill = PatternFill("solid", fgColor="EEF3FB")

    # freeze header row
    ws.freeze_panes = "A2"

    wb.save(OUTPUT_XLS)
    print(f"\n저장 완료 -> {OUTPUT_XLS}")
    print(f"총 {len(rows)}행")


if __name__ == "__main__":
    main()
