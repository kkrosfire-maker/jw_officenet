"""
네이버 블로그 일괄 추출기
────────────────────────────────────────────────────────
입력: 블로그 주소, 연도, 저장 디렉토리
출력:
  posts/{날짜}_{제목}/
      {제목}_이미지_01.jpg  …
      {제목}_텍스트.docx
  {블로그ID}_{연도}_분석.xlsx
      시트1 · 블로그 목록 : 날짜 / 제목 / 요약
      시트2 · 단어 빈도   : 순위 / 단어 / 횟수 / 분류
────────────────────────────────────────────────────────
"""

import os, re, sys, time
from collections import Counter
from datetime import datetime
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from kiwipiepy import Kiwi
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ══════════════════════════════════════════════════════════════════════════════
# 1. 공통 설정
# ══════════════════════════════════════════════════════════════════════════════

SESSION = requests.Session()


def set_headers(blog_id: str):
    SESSION.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Referer": f"https://blog.naver.com/{blog_id}",
        "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8",
    })


def safe_name(text: str, max_len: int = 50) -> str:
    text = re.sub(r'[\\/:*?"<>|]', "", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:max_len]


def ext_from_url(url: str) -> str:
    path = urlparse(url).path
    ext  = os.path.splitext(path)[1].lower()
    return ext if ext in (".jpg", ".jpeg", ".png", ".gif", ".webp") else ".jpg"


def step(n: int, label: str):
    print(f"\n{'─'*60}")
    print(f"  STEP {n} │ {label}")
    print(f"{'─'*60}")


def bar(current: int, total: int, width: int = 30) -> str:
    filled = int(width * current / total)
    return f"[{'█'*filled}{'░'*(width-filled)}] {current}/{total}"


# ══════════════════════════════════════════════════════════════════════════════
# 2. 입력 파싱
# ══════════════════════════════════════════════════════════════════════════════

def parse_blog_id(url: str) -> str:
    """URL 또는 블로그 ID 문자열에서 blog_id 추출."""
    url = url.strip().rstrip("/")
    parsed = urlparse(url)
    if parsed.netloc:                    # URL 형태
        parts = [p for p in parsed.path.split("/") if p]
        if parts:
            return parts[0]
    # 그냥 ID만 입력한 경우
    return url


def get_inputs():
    print("=" * 60)
    print("  BlogMiner — 네이버 블로그 일괄 추출기")
    print("=" * 60)

    raw_url = input("\n블로그 주소 (예: https://blog.naver.com/yesclinic_juan): ").strip()
    blog_id = parse_blog_id(raw_url)
    if not blog_id:
        sys.exit("오류: 블로그 주소를 올바르게 입력해 주세요.")

    year_input = input("추출할 연도 (예: 2023): ").strip()
    try:
        year = int(year_input)
    except ValueError:
        sys.exit("오류: 연도는 숫자로 입력해 주세요.")

    out_default = os.path.join(os.path.expanduser("~"), "Desktop", f"{blog_id}_{year}")
    out_input   = input(f"저장 디렉토리 (기본값 Enter = {out_default}): ").strip()
    out_dir     = out_input if out_input else out_default

    print(f"\n  블로그 ID : {blog_id}")
    print(f"  연도      : {year}")
    print(f"  저장 경로 : {out_dir}")
    confirm = input("\n위 설정으로 시작할까요? (y / n): ").strip().lower()
    if confirm != "y":
        sys.exit("취소되었습니다.")

    return blog_id, year, out_dir


# ══════════════════════════════════════════════════════════════════════════════
# 3. 포스트 목록 수집
# ══════════════════════════════════════════════════════════════════════════════

def get_posts(blog_id: str, year: int) -> list[dict]:
    posts, page, found_older = [], 1, False
    while not found_older and page <= 200:
        try:
            r = SESSION.get(
                f"https://m.blog.naver.com/api/blogs/{blog_id}/post-list"
                f"?categoryNo=0&page={page}",
                timeout=15,
            )
            r.raise_for_status()
            items = r.json()["result"]["items"]
        except Exception as e:
            print(f"  [경고] 페이지 {page} 수신 실패: {e}")
            break
        if not items:
            break
        for item in items:
            dt = datetime.fromtimestamp(item["addDate"] / 1000)
            if dt.year == year:
                posts.append({
                    "logNo": str(item["logNo"]),
                    "title": item.get("titleWithInspectMessage", ""),
                    "date":  dt.strftime("%Y-%m-%d"),
                })
            elif dt.year < year:
                found_older = True
        page += 1
        time.sleep(0.4)
    return posts


# ══════════════════════════════════════════════════════════════════════════════
# 4. 이미지·텍스트 추출 (포스트 1개)
# ══════════════════════════════════════════════════════════════════════════════

SKIP_IMG_DOMAINS = ("nblog", "blogimgs", "ssl.pstatic", "blogpfthumb",
                    "staticmap", ".bin", "tracking", "spc.gif", "btn_")


def clean_img_url(url: str) -> str:
    return url if "postfiles.pstatic.net" in url else url.split("?")[0]


def is_valid_img(url: str, tag) -> bool:
    lower = url.lower()
    if any(x in lower for x in SKIP_IMG_DOMAINS):
        return False
    if not url.startswith("http"):
        return False
    try:
        if tag.get("width")  and int(tag["width"])  < 50: return False
        if tag.get("height") and int(tag["height"]) < 50: return False
    except (ValueError, TypeError):
        pass
    return True


def download_img(url: str, dest: str) -> bool:
    try:
        r = SESSION.get(
            url, timeout=20, stream=True,
            headers={"Accept": "image/avif,image/webp,image/apng,image/*,*/*"},
        )
        r.raise_for_status()
        with open(dest, "wb") as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
        return True
    except Exception as e:
        print(f"    [skip] {url[:60]} - {e}")
        return False


def get_title(soup: BeautifulSoup, log_no: str) -> str:
    for tag_name, attrs in [
        ("h3",  {"class": "se-title-text"}),
        ("div", {"class": "se-title-text"}),
        ("div", {"class": "post-title"}),
        ("h1",  {"class": "htitle"}),
    ]:
        t = soup.find(tag_name, attrs)
        if t:
            return t.get_text(strip=True)
    t = soup.find("title")
    if t:
        return re.sub(r"\s*[:|-]?\s*네이버\s*블로그.*$", "", t.get_text(strip=True)).strip()
    return log_no


def collect_content(body) -> list:
    items, seen_imgs, seen_texts, counter = [], set(), set(), 0
    for elem in body.descendants:
        if not hasattr(elem, "name") or elem.name is None:
            continue
        if elem.name == "img":
            src = elem.get("data-lazy-src") or elem.get("src") or ""
            if not src or src.startswith("data:"):
                continue
            src = clean_img_url(src)
            if src in seen_imgs or not is_valid_img(src, elem):
                continue
            seen_imgs.add(src)
            counter += 1
            items.append(("image", src, counter))
        elif elem.name in ("p","h1","h2","h3","h4","h5","h6","li","span","div"):
            if elem.find(["p","h1","h2","h3","div","li"]) or elem.find("img"):
                continue
            text = elem.get_text(separator=" ", strip=True)
            if not text or len(text) <= 1:
                continue
            key = re.sub(r"\s+", " ", text).strip()
            if key in seen_texts:
                continue
            seen_texts.add(key)
            items.append(("text", text))
    return items


def process_post(blog_id: str, post: dict, posts_dir: str) -> dict:
    log_no = post["logNo"]
    date   = post["date"]
    try:
        resp = SESSION.get(
            f"https://blog.naver.com/PostView.naver"
            f"?blogId={blog_id}&logNo={log_no}&redirect=Dlog&widgetTypeCall=true",
            timeout=20,
        )
        resp.encoding = "utf-8"
    except Exception as e:
        return {"logNo": log_no, "status": "fail", "reason": str(e)}

    soup      = BeautifulSoup(resp.text, "lxml")
    title     = get_title(soup, log_no) or post["title"]
    base_name = safe_name(title)
    folder    = os.path.join(posts_dir, f"{date}_{base_name}")
    os.makedirs(folder, exist_ok=True)

    body = (
        soup.find("div", class_="se-main-container")
        or soup.find("div", id="postViewArea")
        or soup.find("div", class_="post-view")
        or soup.find("div", id="content")
        or soup.body
    )
    content   = collect_content(body or soup)
    img_items = [(item[1], item[2]) for item in content if item[0] == "image"]

    # 이미지 다운로드
    img_files: dict[int, str] = {}
    for url, idx in img_items:
        fname = f"{base_name}_이미지_{idx:02d}{ext_from_url(url)}"
        if download_img(url, os.path.join(folder, fname)):
            img_files[idx] = fname
        time.sleep(0.25)

    # Word 문서 생성
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"원본 URL: https://blog.naver.com/{blog_id}/{log_no}")
    doc.add_paragraph(f"작성일: {date}")
    doc.add_paragraph("")
    text_count = 0
    for item in content:
        if item[0] == "image":
            fname = img_files.get(item[2], f"이미지_{item[2]:02d}")
            p = doc.add_paragraph()
            p.add_run(f"[{os.path.splitext(fname)[0]}]").bold = True
        else:
            p = doc.add_paragraph(item[1])
            p.style.font.size = Pt(11)
            text_count += 1

    docx_name = f"{base_name}_텍스트.docx"
    doc.save(os.path.join(folder, docx_name))

    return {
        "logNo": log_no, "status": "ok",
        "title": title, "date": date,
        "images": len(img_files), "texts": text_count,
        "folder": folder, "docx": os.path.join(folder, docx_name),
    }


# ══════════════════════════════════════════════════════════════════════════════
# 5. 엑셀 생성
# ══════════════════════════════════════════════════════════════════════════════

THIN   = Side(style="thin", color="BFBFBF")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

SKIP_PREFIXES = ("원본 URL:", "작성일:", "http")
SKIP_BRACKET  = re.compile(r"^\[.+\]$")


def hcell(ws, row, col, value, bg="2E75B6", fg="FFFFFF", w=None):
    c = ws.cell(row=row, column=col, value=value)
    c.font      = Font(name="맑은 고딕", bold=True, size=11, color=fg)
    c.fill      = PatternFill("solid", fgColor=bg)
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    c.border    = BORDER
    if w:
        ws.column_dimensions[get_column_letter(col)].width = w
    return c


def dcell(ws, row, col, value, align="left", shade=False, bg_color="EEF3FB"):
    c = ws.cell(row=row, column=col, value=value)
    c.font      = Font(name="맑은 고딕", size=10)
    c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
    c.border    = BORDER
    if shade:
        c.fill = PatternFill("solid", fgColor=bg_color)
    return c


def extract_summary(docx_path: str, heading_text: str, max_chars: int = 120) -> str:
    try:
        doc   = Document(docx_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                continue
            if any(text.startswith(p) for p in SKIP_PREFIXES):
                continue
            if SKIP_BRACKET.match(text):
                continue
            if text == heading_text:
                continue
            parts.append(text)
            if sum(len(p) for p in parts) >= max_chars:
                break
        summary = " ".join(parts)
        if len(summary) > max_chars:
            summary = summary[:max_chars].rsplit(" ", 1)[0] + "…"
        return summary
    except Exception as e:
        return f"(읽기 오류: {e})"


def build_post_sheet(wb, results: list):
    ws = wb.create_sheet(title="블로그 목록")
    ws.row_dimensions[1].height = 28
    hcell(ws, 1, 1, "날짜",   w=13)
    hcell(ws, 1, 2, "제목",   w=42)
    hcell(ws, 1, 3, "요약",   w=72)

    ok = [r for r in results if r["status"] == "ok"]
    ok.sort(key=lambda r: r["date"])
    for i, r in enumerate(ok, 2):
        shade = (i % 2 == 0)
        summary = extract_summary(r["docx"], r["title"])
        ws.row_dimensions[i].height = 40
        dcell(ws, i, 1, r["date"],  "center", shade)
        dcell(ws, i, 2, r["title"], "left",   shade)
        dcell(ws, i, 3, summary,    "left",   shade)

    ws.freeze_panes = "A2"


# ══════════════════════════════════════════════════════════════════════════════
# 6. 단어 빈도 + 분류
# ══════════════════════════════════════════════════════════════════════════════

STOPWORDS = {
    "안녕하세요","있습니다","있는","있어","있고","있으며",
    "합니다","하는","하고","하여","하면","하지","하게","하기",
    "됩니다","되는","되어","되고","되면","되지",
    "이런","이렇게","이러한","이와","이후","이전",
    "그리고","그런","그래서","그러나","그러면","그렇게","그것",
    "때문에","때문","경우","위해","위한","통해","통한",
    "오늘","우리","저희","여러","다양한","중요한","중요",
    "관련","필요","가능","정도","이상","이하","특히",
    "수","것","등","및","의","가","이","을","를",
    "은","는","에","에서","로","으로","도","과","와",
    "이라","라고","이며","이고","지만",
    "많은","많이","또한","함께","통하여","대한","해당",
    "주요","최근","현재","바로","더욱","매우","너무",
    "좋은","좋습니다","좋아","같은","같이",
}

TARGET_POS = {"NNG", "NNP", "VV", "VA"}

MEDICAL_WORDS = {
    "질환","증상","검사","치료","원인","발생","혈압","내시경",
    "예방","검진","관리","진료","식도","건강","위암","역류",
    "당뇨","섭취","대사","발견","대장","증후군","조기","고지혈증",
    "수액","운동","위염","비염","만성","부족","심장","염증",
    "혈관","진단","비만","지방","혈당","기관지염","혈액","초기",
    "알레르기","점막","신체","흡연","내원","스트레스","약물",
    "인슐린","골다공증","당뇨병","통증","두통","발병","식습관",
    "비타민","음주","동맥","감기","축농증","유발","소화","정기",
    "환자","급성","부정맥","면역력","복부","불편","체중",
    "소화기","기능","합병증","뇌졸중","콜레스테롤","천식",
    "초음파","저하","정상","수치","전문의","자극","호흡",
    "소변","신경","가족력","위험","대장암","분비","건조",
    "부비동","감염","조절","위산","회복","의심","경화",
    "쓰리","기침","위장","콧물","바이러스","심근","경색",
    "방치","복용","주사","카페인","가슴","부위",
}

LOCATION_HOSPITAL_WORDS = {
    "연세예스","인천광역시","미추홀구","경인로",
    "주안동내과","미추홀내과","원장","빌딩","국제",
}

CLASSIFY_COLOR = {
    "의료":     "D9EAD3",
    "지역·병원": "CFE2F3",
    "기타":     "F4CCCC",
}


def classify(word: str) -> str:
    if word in MEDICAL_WORDS:
        return "의료"
    if word in LOCATION_HOSPITAL_WORDS:
        return "지역·병원"
    if any(word.endswith(s) for s in ("내과","병원","의원","클리닉","구","시","동","로")):
        return "지역·병원"
    return "기타"


def collect_all_text(results: list) -> str:
    parts = []
    for r in results:
        if r["status"] != "ok":
            continue
        try:
            doc = Document(r["docx"])
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if para.style.name.startswith("Heading"):
                    continue
                if any(text.startswith(p) for p in SKIP_PREFIXES):
                    continue
                if SKIP_BRACKET.match(text):
                    continue
                parts.append(text)
        except Exception:
            pass
    return " ".join(parts)


def build_freq_sheet(wb, results: list, top_n: int = 200):
    print("  한국어 형태소 분석 중...")
    kiwi    = Kiwi()
    text    = collect_all_text(results)
    counter: Counter = Counter()
    chunk   = 5000
    for i in range(0, len(text), chunk):
        for token, pos, *_ in kiwi.analyze(text[i:i+chunk])[0][0]:
            w = token.strip()
            if pos not in TARGET_POS or len(w) < 2 or w in STOPWORDS:
                continue
            if re.fullmatch(r"[0-9]+", w):
                continue
            counter[w] += 1

    top = counter.most_common(top_n)
    print(f"  단어 종류 {len(counter):,}개 중 상위 {len(top)}개 집계 완료")

    ws = wb.create_sheet(title="단어 빈도")
    ws.row_dimensions[1].height = 26
    hcell(ws, 1, 1, "순위", bg="375623", w=8)
    hcell(ws, 1, 2, "단어", bg="375623", w=20)
    hcell(ws, 1, 3, "횟수", bg="375623", w=12)
    hcell(ws, 1, 4, "분류", bg="375623", w=14)

    for rank, (word, count) in enumerate(top, 1):
        r     = rank + 1
        label = classify(word)
        color = CLASSIFY_COLOR[label]
        ws.row_dimensions[r].height = 20
        for col, val, align in [
            (1, rank,  "center"),
            (2, word,  "center"),
            (3, count, "right"),
            (4, label, "center"),
        ]:
            c = ws.cell(row=r, column=col, value=val)
            c.font      = Font(name="맑은 고딕", size=10)
            c.alignment = Alignment(horizontal=align, vertical="center")
            c.border    = BORDER
            c.fill      = PatternFill("solid", fgColor=color)

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:D{len(top)+1}"


# ══════════════════════════════════════════════════════════════════════════════
# 7. 메인
# ══════════════════════════════════════════════════════════════════════════════

def main():
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    blog_id, year, out_dir = get_inputs()
    set_headers(blog_id)

    posts_dir = os.path.join(out_dir, "posts")
    os.makedirs(posts_dir, exist_ok=True)

    # ── STEP 1: 포스트 목록 ────────────────────────────────────────────────────
    step(1, f"{year}년 포스트 목록 수집")
    posts = get_posts(blog_id, year)
    if not posts:
        sys.exit("해당 연도의 포스트를 찾을 수 없습니다.")
    print(f"  → {len(posts)}개 포스트 발견")

    # ── STEP 2: 이미지·텍스트 추출 ────────────────────────────────────────────
    step(2, f"포스트 다운로드 (이미지 + Word 문서)")
    results = []
    for i, post in enumerate(posts, 1):
        print(f"  {bar(i, len(posts))}  {post['date']} {post['title'][:30]}")
        result = process_post(blog_id, post, posts_dir)
        results.append(result)
        time.sleep(0.8)

    ok    = sum(1 for r in results if r["status"] == "ok")
    fails = sum(1 for r in results if r["status"] != "ok")
    print(f"\n  완료: 성공 {ok}개 / 실패 {fails}개")

    # ── STEP 3: 엑셀 생성 ─────────────────────────────────────────────────────
    step(3, "엑셀 파일 생성")
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # 기본 Sheet 제거

    print("  블로그 목록 시트 작성 중...")
    build_post_sheet(wb, results)

    print("  단어 빈도 시트 작성 중...")
    build_freq_sheet(wb, results)

    xlsx_name = f"{blog_id}_{year}_분석.xlsx"
    xlsx_path = os.path.join(out_dir, xlsx_name)
    wb.save(xlsx_path)
    print(f"  저장 완료 → {xlsx_path}")

    # ── 완료 요약 ──────────────────────────────────────────────────────────────
    total_imgs = sum(r.get("images", 0) for r in results if r["status"] == "ok")
    print(f"\n{'='*60}")
    print(f"  완료!")
    print(f"  포스트     : {ok}개")
    print(f"  이미지     : {total_imgs}개")
    print(f"  저장 폴더  : {out_dir}")
    print(f"  엑셀 파일  : {xlsx_name}")
    print(f"{'='*60}")

    # 엑셀 파일 열기
    open_q = input("\n엑셀 파일을 바로 여시겠습니까? (y / n): ").strip().lower()
    if open_q == "y":
        os.startfile(xlsx_path)


if __name__ == "__main__":
    main()
