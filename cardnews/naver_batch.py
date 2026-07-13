"""
Naver blog batch extractor — processes all 2023 posts from a given blog.
Each post gets its own subfolder:  output/<date>_<title>/
"""
import os
import re
import sys
import time
import requests
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from urllib.parse import urlparse

# ── config ─────────────────────────────────────────────────────────────────────
BLOG_ID     = "yesclinic_juan"
TARGET_YEAR = 2023
OUTPUT_ROOT = os.path.join(os.path.dirname(__file__), "naver_output_2023")

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Referer": f"https://blog.naver.com/{BLOG_ID}",
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8",
}
SESSION = requests.Session()
SESSION.headers.update(HEADERS)

# ── helpers ────────────────────────────────────────────────────────────────────

def safe_filename(text: str, max_len: int = 50) -> str:
    text = re.sub(r'[\\/:*?"<>|]', "", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:max_len]


def fetch(url: str, **kwargs) -> requests.Response:
    resp = SESSION.get(url, timeout=20, **kwargs)
    resp.raise_for_status()
    return resp


def ext_from_url(url: str) -> str:
    path = urlparse(url).path
    ext = os.path.splitext(path)[1].lower()
    return ext if ext in (".jpg", ".jpeg", ".png", ".gif", ".webp") else ".jpg"


def clean_img_url(url: str) -> str:
    if "postfiles.pstatic.net" in url:
        return url  # keep ?type=w966 — required for access
    return url.split("?")[0]


def is_valid_img(url: str, tag) -> bool:
    lower = url.lower()
    if any(x in lower for x in ("staticmap", ".bin", "tracking", "icon", "btn_",
                                  "spc.gif", "nblog", "blogimgs", "ssl.pstatic",
                                  "blogpfthumb")):
        return False
    if not url.startswith("http"):
        return False
    w, h = tag.get("width", ""), tag.get("height", "")
    try:
        if w and int(w) < 50: return False
        if h and int(h) < 50: return False
    except ValueError:
        pass
    return True


def download_image(url: str, dest: str) -> bool:
    try:
        r = SESSION.get(url, timeout=20, stream=True,
                        headers={"Accept": "image/avif,image/webp,image/apng,image/*,*/*"})
        r.raise_for_status()
        with open(dest, "wb") as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
        return True
    except Exception as e:
        print(f"    [skip] {url[:60]} - {e}")
        return False


def get_title(soup: BeautifulSoup, log_no: str) -> str:
    for sel in [
        ("h3", {"class": "se-title-text"}),
        ("div", {"class": "se-title-text"}),
        ("div", {"class": "post-title"}),
        ("h1", {"class": "htitle"}),
    ]:
        tag = soup.find(*sel)
        if tag:
            return tag.get_text(strip=True)
    tag = soup.find("title")
    if tag:
        text = re.sub(r"\s*[:|-]?\s*네이버\s*블로그.*$", "", tag.get_text(strip=True)).strip()
        if text:
            return text
    return log_no


def collect_content(body) -> list:
    """Return ordered list of ('image', url, idx) and ('text', text) items."""
    items, seen_imgs, seen_texts, counter = [], set(), set(), 0

    for elem in body.descendants:
        if not hasattr(elem, "name") or elem.name is None:
            continue

        if elem.name == "img":
            src = elem.get("data-lazy-src") or elem.get("src") or ""
            if not src or src.startswith("data:"):
                continue
            src = clean_img_url(src)
            if src in seen_imgs or not is_valid_img(src, elem):
                continue
            seen_imgs.add(src)
            counter += 1
            items.append(("image", src, counter))

        elif elem.name in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "span", "div"):
            if elem.find(["p", "h1", "h2", "h3", "div", "li"]):
                continue
            if elem.find("img"):
                continue
            text = elem.get_text(separator=" ", strip=True)
            if not text or len(text) <= 1:
                continue
            key = re.sub(r"\s+", " ", text).strip()
            if key in seen_texts:
                continue
            seen_texts.add(key)
            items.append(("text", text))

    return items


# ── step 1: get list of all 2023 posts ────────────────────────────────────────

def get_2023_posts() -> list[dict]:
    posts = []
    page = 1
    found_older = False

    print(f"[1/3] 2023년 포스트 목록 수집 중...")
    while not found_older and page <= 200:
        try:
            r = fetch(
                f"https://m.blog.naver.com/api/blogs/{BLOG_ID}/post-list"
                f"?categoryNo=0&page={page}"
            )
            items = r.json()["result"]["items"]
        except Exception as e:
            print(f"  페이지 {page} 오류: {e}")
            break

        if not items:
            break

        for item in items:
            dt = datetime.fromtimestamp(item["addDate"] / 1000)
            if dt.year == TARGET_YEAR:
                posts.append({
                    "logNo": str(item["logNo"]),
                    "title": item.get("titleWithInspectMessage", ""),
                    "date":  dt.strftime("%Y-%m-%d"),
                })
            elif dt.year < TARGET_YEAR:
                found_older = True

        page += 1
        time.sleep(0.5)

    print(f"  -> {len(posts)}개 포스트 발견")
    return posts


# ── step 2: process one post ───────────────────────────────────────────────────

def process_post(post: dict, index: int, total: int) -> dict:
    log_no = post["logNo"]
    date   = post["date"]
    print(f"\n[{index}/{total}] {date} | {post['title'][:40]}")

    # fetch post HTML (iframe content)
    try:
        resp = fetch(
            f"https://blog.naver.com/PostView.naver"
            f"?blogId={BLOG_ID}&logNo={log_no}&redirect=Dlog&widgetTypeCall=true"
        )
        resp.encoding = "utf-8"
        html = resp.text
    except Exception as e:
        print(f"  [오류] HTML 수신 실패: {e}")
        return {"logNo": log_no, "status": "fail", "reason": str(e)}

    soup = BeautifulSoup(html, "lxml")

    title     = get_title(soup, log_no) or post["title"]
    base_name = safe_filename(title)
    folder    = os.path.join(OUTPUT_ROOT, f"{date}_{base_name}")
    os.makedirs(folder, exist_ok=True)

    body = (
        soup.find("div", class_="se-main-container")
        or soup.find("div", id="postViewArea")
        or soup.find("div", class_="post-view")
        or soup.find("div", id="content")
        or soup.body
    )

    content   = collect_content(body or soup)
    img_items = [(item[1], item[2]) for item in content if item[0] == "image"]
    print(f"  이미지 {len(img_items)}개, 텍스트 단락 {sum(1 for i in content if i[0]=='text')}개")

    # download images
    img_files: dict[int, str] = {}
    for url, idx in img_items:
        ext   = ext_from_url(url)
        fname = f"{base_name}_이미지_{idx:02d}{ext}"
        dest  = os.path.join(folder, fname)
        if download_image(url, dest):
            img_files[idx] = fname
        time.sleep(0.3)

    # build Word document
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"원본 URL: https://blog.naver.com/{BLOG_ID}/{log_no}")
    doc.add_paragraph(f"작성일: {date}")
    doc.add_paragraph("")

    text_count = 0
    for item in content:
        if item[0] == "image":
            _, url, idx = item
            fname = img_files.get(idx, f"이미지_{idx:02d}")
            label = os.path.splitext(fname)[0]
            p = doc.add_paragraph()
            p.add_run(f"[{label}]").bold = True
        else:
            _, text = item
            p = doc.add_paragraph(text)
            p.style.font.size = Pt(11)
            text_count += 1

    docx_name = f"{base_name}_텍스트.docx"
    doc.save(os.path.join(folder, docx_name))

    print(f"  저장 완료 -> {folder}")
    return {
        "logNo":  log_no,
        "status": "ok",
        "images": len(img_files),
        "texts":  text_count,
        "folder": folder,
    }


# ── main ───────────────────────────────────────────────────────────────────────

def main():
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    os.makedirs(OUTPUT_ROOT, exist_ok=True)

    posts = get_2023_posts()
    if not posts:
        print("포스트를 찾을 수 없습니다.")
        return

    print(f"\n[2/3] 포스트 처리 시작 (총 {len(posts)}개)")
    results = []
    for i, post in enumerate(posts, 1):
        result = process_post(post, i, len(posts))
        results.append(result)
        time.sleep(1.0)  # be polite to Naver servers

    # summary
    ok    = [r for r in results if r["status"] == "ok"]
    fails = [r for r in results if r["status"] != "ok"]

    print(f"\n[3/3] 완료 요약")
    print(f"  성공: {len(ok)}개 / 실패: {len(fails)}개")
    print(f"  저장 폴더: {OUTPUT_ROOT}")
    if fails:
        print("  실패 포스트:")
        for f in fails:
            print(f"    logNo={f['logNo']} - {f.get('reason','')}")


if __name__ == "__main__":
    main()
