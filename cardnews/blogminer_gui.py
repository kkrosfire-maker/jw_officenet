"""
BlogMiner GUI — 네이버 블로그 일괄 추출기
탭 1: 추출   탭 2: 결과 미리보기 + 단어 치환
"""
import os, re, sys, time, threading, queue
from collections import Counter
from datetime import datetime
from urllib.parse import urlparse
from tkinter import *
from tkinter import ttk, filedialog, messagebox

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from kiwipiepy import Kiwi
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ══════════════════════════════════════════════════════════════════
# 색상 / 폰트
# ══════════════════════════════════════════════════════════════════
C = {
    "bg":      "#F7F9FC",
    "header":  "#1A365D",
    "accent":  "#2B6CB0",
    "success": "#276749",
    "panel":   "#FFFFFF",
    "border":  "#CBD5E0",
    "txt":     "#2D3748",
    "txt2":    "#718096",
    "log_bg":  "#1A202C",
    "log_txt": "#E2E8F0",
    "row_alt": "#EBF8FF",
}
FONT      = ("맑은 고딕", 10)
FONT_BOLD = ("맑은 고딕", 10, "bold")
FONT_SM   = ("맑은 고딕",  9)

# ══════════════════════════════════════════════════════════════════
# 추출 핵심 로직
# ══════════════════════════════════════════════════════════════════
SESSION = requests.Session()

def set_session_headers(blog_id):
    SESSION.headers.update({
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                       "AppleWebKit/537.36 Chrome/124.0.0.0 Safari/537.36"),
        "Referer": f"https://blog.naver.com/{blog_id}",
        "Accept-Language": "ko-KR,ko;q=0.9",
    })

def parse_blog_id(url):
    url = url.strip().rstrip("/")
    parsed = urlparse(url)
    if parsed.netloc:
        parts = [p for p in parsed.path.split("/") if p]
        if parts: return parts[0]
    return url

def safe_name(text, max_len=50):
    text = re.sub(r'[\\/:*?"<>|]', "", text)
    text = re.sub(r"\s+", "_", text.strip())
    return text[:max_len]

def ext_from_url(url):
    ext = os.path.splitext(urlparse(url).path)[1].lower()
    return ext if ext in (".jpg",".jpeg",".png",".gif",".webp") else ".jpg"

SKIP_IMG = ("nblog","blogimgs","ssl.pstatic","blogpfthumb",
            "staticmap",".bin","tracking","spc.gif","btn_")

def clean_img_url(url):
    return url if "postfiles.pstatic.net" in url else url.split("?")[0]

def is_valid_img(url, tag):
    if any(x in url.lower() for x in SKIP_IMG): return False
    if not url.startswith("http"): return False
    try:
        if tag.get("width")  and int(tag["width"])  < 50: return False
        if tag.get("height") and int(tag["height"]) < 50: return False
    except: pass
    return True

def download_img(url, dest, log):
    try:
        r = SESSION.get(url, timeout=20, stream=True,
                        headers={"Accept":"image/avif,image/webp,image/apng,image/*,*/*"})
        r.raise_for_status()
        with open(dest,"wb") as f:
            for chunk in r.iter_content(8192): f.write(chunk)
        return True
    except Exception as e:
        log(f"    [skip] {url[:55]}", "warn"); return False

def get_title(soup, log_no):
    for tag, attrs in [("h3",{"class":"se-title-text"}),("div",{"class":"se-title-text"}),
                        ("div",{"class":"post-title"}),("h1",{"class":"htitle"})]:
        t = soup.find(tag, attrs)
        if t: return t.get_text(strip=True)
    t = soup.find("title")
    if t: return re.sub(r"\s*[:|-]?\s*네이버\s*블로그.*$","",t.get_text(strip=True)).strip()
    return log_no

def collect_content(body):
    items, seen_imgs, seen_texts, counter = [], set(), set(), 0
    for elem in body.descendants:
        if not hasattr(elem,"name") or not elem.name: continue
        if elem.name == "img":
            src = elem.get("data-lazy-src") or elem.get("src") or ""
            if not src or src.startswith("data:"): continue
            src = clean_img_url(src)
            if src in seen_imgs or not is_valid_img(src,elem): continue
            seen_imgs.add(src); counter += 1
            items.append(("image",src,counter))
        elif elem.name in ("p","h1","h2","h3","h4","h5","h6","li","span","div"):
            if elem.find(["p","h1","h2","h3","div","li"]) or elem.find("img"): continue
            text = elem.get_text(separator=" ",strip=True)
            if not text or len(text)<=1: continue
            key = re.sub(r"\s+"," ",text).strip()
            if key in seen_texts: continue
            seen_texts.add(key); items.append(("text",text))
    return items

def get_posts(blog_id, year, month, log):
    posts, page, found_older = [], 1, False
    while not found_older and page<=200:
        try:
            r = SESSION.get(f"https://m.blog.naver.com/api/blogs/{blog_id}/post-list"
                            f"?categoryNo=0&page={page}", timeout=15)
            r.raise_for_status(); items = r.json()["result"]["items"]
        except Exception as e:
            log(f"  [경고] 페이지 {page} 오류: {e}","warn"); break
        if not items: break
        for item in items:
            dt = datetime.fromtimestamp(item["addDate"]/1000)
            if dt.year == year:
                if month is None or dt.month == month:
                    posts.append({"logNo":str(item["logNo"]),
                                   "title":item.get("titleWithInspectMessage",""),
                                   "date":dt.strftime("%Y-%m-%d")})
            elif dt.year < year: found_older = True
        page+=1; time.sleep(0.4)
    return posts

def process_post(blog_id, post, posts_dir, log):
    log_no=post["logNo"]; date=post["date"]
    try:
        resp = SESSION.get(
            f"https://blog.naver.com/PostView.naver"
            f"?blogId={blog_id}&logNo={log_no}&redirect=Dlog&widgetTypeCall=true",
            timeout=20)
        resp.encoding="utf-8"
    except Exception as e:
        return {"logNo":log_no,"status":"fail","reason":str(e)}
    soup      = BeautifulSoup(resp.text,"lxml")
    title     = get_title(soup,log_no) or post["title"]
    base_name = safe_name(title)
    folder    = os.path.join(posts_dir,f"{date}_{base_name}")
    os.makedirs(folder,exist_ok=True)
    body = (soup.find("div",class_="se-main-container")
            or soup.find("div",id="postViewArea")
            or soup.find("div",class_="post-view")
            or soup.find("div",id="content") or soup.body)
    content   = collect_content(body or soup)
    img_items = [(item[1],item[2]) for item in content if item[0]=="image"]
    img_files = {}
    for url,idx in img_items:
        fname=f"{base_name}_이미지_{idx:02d}{ext_from_url(url)}"
        if download_img(url,os.path.join(folder,fname),log): img_files[idx]=fname
        time.sleep(0.25)
    doc = Document()
    doc.add_heading(title,level=1)
    doc.add_paragraph(f"원본 URL: https://blog.naver.com/{blog_id}/{log_no}")
    doc.add_paragraph(f"작성일: {date}"); doc.add_paragraph("")
    text_count=0
    for item in content:
        if item[0]=="image":
            fname=img_files.get(item[2],f"이미지_{item[2]:02d}")
            p=doc.add_paragraph(); p.add_run(f"[{os.path.splitext(fname)[0]}]").bold=True
        else:
            p=doc.add_paragraph(item[1]); p.style.font.size=Pt(11); text_count+=1
    docx_name=f"{base_name}_텍스트.docx"
    doc.save(os.path.join(folder,docx_name))
    return {"logNo":log_no,"status":"ok","title":title,"date":date,
            "images":len(img_files),"texts":text_count,"folder":folder,
            "docx":os.path.join(folder,docx_name)}

# Excel
THIN=Side(style="thin",color="BFBFBF")
BORDER=Border(left=THIN,right=THIN,top=THIN,bottom=THIN)
SKIP_P=("원본 URL:","작성일:","http"); SKIP_B=re.compile(r"^\[.+\]$")

def _hcell(ws,r,c,v,bg="2E75B6",fg="FFFFFF",w=None):
    cell=ws.cell(row=r,column=c,value=v)
    cell.font=Font(name="맑은 고딕",bold=True,size=11,color=fg)
    cell.fill=PatternFill("solid",fgColor=bg)
    cell.alignment=Alignment(horizontal="center",vertical="center",wrap_text=True)
    cell.border=BORDER
    if w: ws.column_dimensions[get_column_letter(c)].width=w

def _dcell(ws,r,c,v,align="left",shade=False,color="EEF3FB"):
    cell=ws.cell(row=r,column=c,value=v)
    cell.font=Font(name="맑은 고딕",size=10)
    cell.alignment=Alignment(horizontal=align,vertical="center",wrap_text=True)
    cell.border=BORDER
    if shade: cell.fill=PatternFill("solid",fgColor=color)

def extract_summary(docx_path, heading, max_c=120):
    try:
        doc=Document(docx_path); parts=[]
        for para in doc.paragraphs:
            t=para.text.strip()
            if not t or para.style.name.startswith("Heading"): continue
            if any(t.startswith(p) for p in SKIP_P) or SKIP_B.match(t) or t==heading: continue
            parts.append(t)
            if sum(len(p) for p in parts)>=max_c: break
        s=" ".join(parts)
        return (s[:max_c].rsplit(" ",1)[0]+"…") if len(s)>max_c else s
    except: return ""

def build_post_sheet(wb, results):
    ws=wb.create_sheet(title="블로그 목록"); ws.row_dimensions[1].height=28
    _hcell(ws,1,1,"날짜",w=13); _hcell(ws,1,2,"제목",w=42); _hcell(ws,1,3,"요약",w=72)
    ok=sorted([r for r in results if r["status"]=="ok"],key=lambda r:r["date"])
    for i,r in enumerate(ok,2):
        shade=(i%2==0); s=extract_summary(r["docx"],r["title"])
        ws.row_dimensions[i].height=40
        _dcell(ws,i,1,r["date"],"center",shade); _dcell(ws,i,2,r["title"],"left",shade)
        _dcell(ws,i,3,s,"left",shade)
    ws.freeze_panes="A2"

STOPWORDS={"안녕하세요","있습니다","있는","있어","있고","있으며","합니다","하는","하고",
           "하여","하면","하지","하게","하기","됩니다","되는","되어","되고","되면","되지",
           "이런","이렇게","이러한","이와","이후","이전","그리고","그런","그래서","그러나",
           "그러면","그렇게","그것","때문에","때문","경우","위해","위한","통해","통한",
           "오늘","우리","저희","여러","다양한","중요한","중요","관련","필요","가능",
           "정도","이상","이하","특히","수","것","등","및","의","가","이","을","를",
           "은","는","에","에서","로","으로","도","과","와","이라","라고","이며","이고",
           "지만","많은","많이","또한","함께","통하여","대한","해당","주요","최근","현재",
           "바로","더욱","매우","너무","좋은","좋습니다","좋아","같은","같이"}
TARGET_POS={"NNG","NNP","VV","VA"}
MEDICAL={"질환","증상","검사","치료","원인","발생","혈압","내시경","예방","검진","관리",
         "진료","식도","건강","위암","역류","당뇨","섭취","대사","발견","대장","증후군",
         "조기","고지혈증","수액","운동","위염","비염","만성","부족","심장","염증",
         "혈관","진단","비만","지방","혈당","기관지염","혈액","초기","알레르기","점막",
         "신체","흡연","내원","스트레스","약물","인슐린","골다공증","당뇨병","통증",
         "두통","발병","식습관","비타민","음주","동맥","감기","축농증","유발","소화",
         "정기","환자","급성","부정맥","면역력","복부","불편","체중","소화기","기능",
         "합병증","뇌졸중","콜레스테롤","천식","초음파","저하","정상","수치","전문의",
         "자극","호흡","소변","신경","가족력","위험","대장암","분비","건조","부비동",
         "감염","조절","위산","회복","의심","경화","쓰리","기침","위장","콧물",
         "바이러스","심근","경색","방치","복용","주사","카페인","가슴","부위"}
LOC_HOSP={"연세예스","인천광역시","미추홀구","경인로","주안동내과","미추홀내과","원장","빌딩","국제"}
CLF_COLOR={"의료":"D9EAD3","지역·병원":"CFE2F3","기타":"F4CCCC"}

def classify(word):
    if word in MEDICAL: return "의료"
    if word in LOC_HOSP: return "지역·병원"
    if any(word.endswith(s) for s in ("내과","병원","의원","클리닉","구","시","동","로")): return "지역·병원"
    return "기타"

def build_freq_sheet(wb, results, log):
    log("  한국어 형태소 분석 중 (잠시 기다려 주세요)…")
    kiwi=Kiwi(); parts=[]
    for r in results:
        if r["status"]!="ok": continue
        try:
            doc=Document(r["docx"])
            for para in doc.paragraphs:
                t=para.text.strip()
                if not t or para.style.name.startswith("Heading"): continue
                if any(t.startswith(p) for p in SKIP_P) or SKIP_B.match(t): continue
                parts.append(t)
        except: pass
    text=" ".join(parts); counter: Counter=Counter()
    for i in range(0,len(text),5000):
        for token,pos,*_ in kiwi.analyze(text[i:i+5000])[0][0]:
            w=token.strip()
            if pos not in TARGET_POS or len(w)<2 or w in STOPWORDS: continue
            if re.fullmatch(r"[0-9]+",w): continue
            counter[w]+=1
    top=counter.most_common(200)
    log(f"  단어 {len(counter):,}종 → 상위 {len(top)}개 저장")
    ws=wb.create_sheet(title="단어 빈도"); ws.row_dimensions[1].height=26
    _hcell(ws,1,1,"순위",bg="375623",w=8); _hcell(ws,1,2,"단어",bg="375623",w=20)
    _hcell(ws,1,3,"횟수",bg="375623",w=12); _hcell(ws,1,4,"분류",bg="375623",w=14)
    for rank,(word,count) in enumerate(top,1):
        r=rank+1; label=classify(word); color=CLF_COLOR[label]
        ws.row_dimensions[r].height=20
        for c,v,a in [(1,rank,"center"),(2,word,"center"),(3,count,"right"),(4,label,"center")]:
            cell=ws.cell(row=r,column=c,value=v)
            cell.font=Font(name="맑은 고딕",size=10)
            cell.alignment=Alignment(horizontal=a,vertical="center")
            cell.border=BORDER; cell.fill=PatternFill("solid",fgColor=color)
    ws.freeze_panes="A2"; ws.auto_filter.ref=f"A1:D{len(top)+1}"
    return top  # 반환해서 GUI 미리보기에 사용

# ══════════════════════════════════════════════════════════════════
# GUI
# ══════════════════════════════════════════════════════════════════
class BlogMinerApp:
    def __init__(self, root: Tk):
        self.root=root
        self.root.title("BlogMiner v0.1")
        self.root.geometry("760x740")
        self.root.minsize(700,620)
        self.root.configure(bg=C["bg"])
        self._q: queue.Queue=queue.Queue()
        self._running=False
        self._out_dir=""
        self._posts_dir=""
        self._xlsx_path=""
        self._top_words=[]
        self._setup_style()
        self._build_ui()

    # ── 스타일 ─────────────────────────────────────────────────────
    def _setup_style(self):
        s=ttk.Style(self.root); s.theme_use("clam")
        s.configure("TFrame",       background=C["bg"])
        s.configure("Panel.TFrame", background=C["panel"])
        s.configure("TLabel",       background=C["bg"],    foreground=C["txt"], font=FONT)
        s.configure("Panel.TLabel", background=C["panel"], foreground=C["txt"], font=FONT)
        s.configure("Bold.TLabel",  background=C["panel"], foreground=C["txt"], font=FONT_BOLD)
        s.configure("Hint.TLabel",  background=C["panel"], foreground=C["txt2"],font=FONT_SM)
        s.configure("TEntry",       font=FONT, fieldbackground="white", relief="flat")
        s.configure("TNotebook",    background=C["bg"], tabmargins=[2,4,0,0])
        s.configure("TNotebook.Tab",background=C["border"],foreground=C["txt2"],
                    font=FONT_BOLD, padding=[14,6])
        s.map("TNotebook.Tab",
              background=[("selected",C["panel"])],
              foreground=[("selected",C["accent"])])
        s.configure("Accent.TButton",background=C["accent"],foreground="white",
                    font=("맑은 고딕",11,"bold"),padding=(18,7),relief="flat")
        s.map("Accent.TButton",
              background=[("active","#2C5282"),("disabled","#A0AEC0")],
              foreground=[("disabled","#E2E8F0")])
        s.configure("Sub.TButton",background=C["border"],foreground=C["txt"],
                    font=FONT,padding=(10,5),relief="flat")
        s.map("Sub.TButton",background=[("active","#BEE3F8")])
        s.configure("Green.Horizontal.TProgressbar",
                    troughcolor="#E2E8F0",background="#48BB78",thickness=10)
        s.configure("Blue.Horizontal.TProgressbar",
                    troughcolor="#E2E8F0",background=C["accent"],thickness=10)

    # ── 전체 UI ────────────────────────────────────────────────────
    def _build_ui(self):
        # 헤더
        hdr=Frame(self.root,bg=C["header"],height=60)
        hdr.pack(fill=X); hdr.pack_propagate(False)
        Label(hdr,text="BlogMiner",bg=C["header"],fg="white",
              font=("맑은 고딕",17,"bold")).pack(side=LEFT,padx=(20,6),pady=10)
        Label(hdr,text="v0.1",bg=C["header"],fg="#90CDF4",
              font=("맑은 고딕",10,"bold")).pack(side=LEFT,pady=17)
        Label(hdr,text="  네이버 블로그 일괄 추출기",bg=C["header"],fg="#90CDF4",
              font=("맑은 고딕",10)).pack(side=LEFT,pady=17)

        # 탭
        self.nb=ttk.Notebook(self.root)
        self.nb.pack(fill=BOTH,expand=True,padx=14,pady=12)
        self._build_tab_extract()
        self._build_tab_tools()

    # ══════════════════════════════════════════════════════════════
    # TAB 1 — 추출
    # ══════════════════════════════════════════════════════════════
    def _build_tab_extract(self):
        tab=Frame(self.nb,bg=C["bg"]); self.nb.add(tab,text="  📥  추출  ")

        # 입력 패널
        inp=Frame(tab,bg=C["panel"],highlightbackground=C["border"],highlightthickness=1)
        inp.pack(fill=X,padx=4,pady=(6,8))
        Label(inp,text="블로그 정보 입력",bg=C["panel"],fg=C["accent"],
              font=FONT_BOLD).grid(row=0,column=0,columnspan=3,sticky=W,padx=16,pady=(12,6))

        for row_i,(label,hint) in enumerate([
            ("블로그 주소","예: https://blog.naver.com/yesclinic_juan"),
            ("연도 / 월","연도 필수  ·  월은 선택 (특정 월만 추출할 때 입력)"),
            ("저장 폴더",""),
        ],1):
            Label(inp,text=label,bg=C["panel"],fg=C["txt"],
                  font=FONT_BOLD).grid(row=row_i*2-1,column=0,sticky=W,padx=16,pady=(8,2))
            if hint:
                Label(inp,text=hint,bg=C["panel"],fg=C["txt2"],
                      font=FONT_SM).grid(row=row_i*2,column=0,columnspan=3,sticky=W,padx=16,pady=(0,4))

        self.url_var  =StringVar()
        self.year_var =StringVar(value=str(datetime.now().year-1))
        self.month_var=StringVar(value="전체")
        self.dir_var  =StringVar(value=os.path.join(os.path.expanduser("~"),"Desktop"))

        ttk.Entry(inp,textvariable=self.url_var,width=54).grid(
            row=1,column=1,columnspan=2,sticky=EW,padx=(0,16),pady=(8,2))

        yr_f=Frame(inp,bg=C["panel"]); yr_f.grid(row=3,column=1,columnspan=2,sticky=W,pady=(8,4))
        ttk.Entry(yr_f,textvariable=self.year_var,width=8).pack(side=LEFT)
        Label(yr_f,text="년",bg=C["panel"],fg=C["txt"],font=FONT).pack(side=LEFT,padx=(4,10))
        ttk.Combobox(yr_f,textvariable=self.month_var,width=5,state="readonly",
                     values=["전체","1","2","3","4","5","6",
                             "7","8","9","10","11","12"]).pack(side=LEFT)
        Label(yr_f,text="월  (전체 = 연간 전체 추출)",
              bg=C["panel"],fg=C["txt2"],font=FONT_SM).pack(side=LEFT,padx=(6,0))
        dir_e=ttk.Entry(inp,textvariable=self.dir_var,width=42)
        dir_e.grid(row=5,column=1,sticky=EW,padx=(0,4),pady=(8,12))
        ttk.Button(inp,text="찾아보기",style="Sub.TButton",
                   command=self._browse).grid(row=5,column=2,padx=(0,16),pady=(8,12))
        inp.columnconfigure(1,weight=1)

        # 시작 버튼
        bf=Frame(tab,bg=C["bg"]); bf.pack(fill=X,padx=4,pady=(0,8))
        self.start_btn=ttk.Button(bf,text="▶  추출 시작",style="Accent.TButton",
                                  command=self._start)
        self.start_btn.pack(side=RIGHT)

        # 진행 패널
        prog=Frame(tab,bg=C["panel"],highlightbackground=C["border"],highlightthickness=1)
        prog.pack(fill=BOTH,expand=True,padx=4)
        Label(prog,text="진행 상황",bg=C["panel"],fg=C["accent"],
              font=FONT_BOLD).pack(anchor=W,padx=16,pady=(12,6))

        self.step_data=[]
        for label,desc in [("STEP 1","포스트 목록 수집"),
                            ("STEP 2","이미지 · 문서 다운로드"),
                            ("STEP 3","엑셀 파일 생성")]:
            row=Frame(prog,bg=C["panel"]); row.pack(fill=X,padx=16,pady=3)
            Label(row,text=label,bg=C["panel"],fg=C["accent"],
                  font=("맑은 고딕",9,"bold"),width=7).pack(side=LEFT)
            Label(row,text=desc,bg=C["panel"],fg=C["txt"],
                  font=FONT_SM,width=20,anchor=W).pack(side=LEFT)
            pb=ttk.Progressbar(row,style="Blue.Horizontal.TProgressbar",
                               length=280,mode="determinate")
            pb.pack(side=LEFT,padx=8)
            stat=Label(row,text="대기",bg=C["panel"],fg=C["txt2"],
                       font=FONT_SM,width=14,anchor=W)
            stat.pack(side=LEFT)
            self.step_data.append((pb,stat))

        Frame(prog,bg=C["border"],height=1).pack(fill=X,padx=16,pady=(8,0))
        lf=Frame(prog,bg=C["panel"]); lf.pack(fill=BOTH,expand=True,padx=16,pady=(6,16))
        self.log_text=Text(lf,bg=C["log_bg"],fg=C["log_txt"],font=("Consolas",9),
                           relief="flat",wrap=WORD,state=DISABLED,height=10)
        sb=ttk.Scrollbar(lf,orient=VERTICAL,command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=sb.set)
        self.log_text.pack(side=LEFT,fill=BOTH,expand=True)
        sb.pack(side=RIGHT,fill=Y)
        self.log_text.tag_config("step",foreground="#68D391")
        self.log_text.tag_config("ok",  foreground="#68D391")
        self.log_text.tag_config("warn",foreground="#F6AD55")
        self.log_text.tag_config("err", foreground="#FC8181")
        self.log_text.tag_config("plain",foreground=C["log_txt"])

    # ══════════════════════════════════════════════════════════════
    # TAB 2 — 단어 빈도 + 인라인 치환 테이블
    # ══════════════════════════════════════════════════════════════
    def _build_tab_tools(self):
        tab=Frame(self.nb,bg=C["bg"]); self.nb.add(tab,text="  🔤  결과 · 단어 치환  ")
        self._replace_rows=[]

        # ── 상단 버튼 바 ────────────────────────────────────────────
        top_bar=Frame(tab,bg=C["bg"]); top_bar.pack(fill=X,padx=4,pady=(6,2))
        Label(top_bar,text="단어 빈도 · 치환 (상위 200개)",bg=C["bg"],
              fg=C["accent"],font=FONT_BOLD).pack(side=LEFT)
        self.open_xlsx_btn=ttk.Button(top_bar,text="📊 엑셀 파일 열기",style="Sub.TButton",
                                      command=self._open_xlsx,state=DISABLED)
        self.open_xlsx_btn.pack(side=RIGHT)
        ttk.Button(top_bar,text="✏  치환 적용",style="Accent.TButton",
                   command=self._do_replace).pack(side=RIGHT,padx=(0,6))

        Label(tab,text="'바꿀 단어' 칸을 입력한 항목만 치환됩니다. 비워두면 건너뜁니다.",
              bg=C["bg"],fg=C["txt2"],font=FONT_SM).pack(anchor=W,padx=6,pady=(0,4))

        # ── 작업 폴더 ──────────────────────────────────────────────
        fp=Frame(tab,bg=C["panel"],highlightbackground=C["border"],highlightthickness=1)
        fp.pack(fill=X,padx=4,pady=(0,4))
        Label(fp,text="작업 폴더",bg=C["panel"],fg=C["txt"],
              font=FONT_BOLD).pack(side=LEFT,padx=(16,8),pady=8)
        self.rep_dir_var=StringVar()
        ttk.Entry(fp,textvariable=self.rep_dir_var).pack(
            side=LEFT,fill=X,expand=True,padx=(0,4),pady=8)
        ttk.Button(fp,text="찾아보기",style="Sub.TButton",
                   command=self._browse_rep).pack(side=LEFT,padx=(0,16),pady=8)

        # ── 단어 테이블 (스크롤 가능 캔버스) ──────────────────────
        outer=Frame(tab,bg=C["panel"],highlightbackground=C["border"],highlightthickness=1)
        outer.pack(fill=BOTH,expand=True,padx=4,pady=(0,4))

        hdr=Frame(outer,bg=C["accent"]); hdr.pack(fill=X)
        for txt,w in [("순위",5),("단어",14),("횟수",7),("분류",10),("바꿀 단어",18)]:
            Label(hdr,text=txt,bg=C["accent"],fg="white",font=FONT_BOLD,
                  width=w,anchor="center",pady=5).pack(side=LEFT,padx=1)

        cf=Frame(outer,bg=C["panel"]); cf.pack(fill=BOTH,expand=True)
        self._repl_canvas=Canvas(cf,bg=C["panel"],highlightthickness=0)
        _sb=ttk.Scrollbar(cf,orient=VERTICAL,command=self._repl_canvas.yview)
        self._repl_canvas.configure(yscrollcommand=_sb.set)
        _sb.pack(side=RIGHT,fill=Y)
        self._repl_canvas.pack(side=LEFT,fill=BOTH,expand=True)

        self._repl_inner=Frame(self._repl_canvas,bg=C["panel"])
        self._repl_win=self._repl_canvas.create_window(
            (0,0),window=self._repl_inner,anchor="nw")
        self._repl_inner.bind("<Configure>",
            lambda e: self._repl_canvas.configure(
                scrollregion=self._repl_canvas.bbox("all")))
        self._repl_canvas.bind("<Configure>",
            lambda e: self._repl_canvas.itemconfig(self._repl_win,width=e.width))
        self._repl_canvas.bind("<MouseWheel>",self._on_repl_scroll)

        Label(self._repl_inner,
              text="추출이 완료되면 여기에 단어 목록이 표시됩니다.",
              bg=C["panel"],fg=C["txt2"],font=FONT_SM).pack(pady=20)

        # ── 치환 결과 로그 ─────────────────────────────────────────
        rl=Frame(tab,bg=C["panel"],highlightbackground=C["border"],highlightthickness=1)
        rl.pack(fill=X,padx=4,pady=(0,4))
        Label(rl,text="치환 결과",bg=C["panel"],fg=C["accent"],
              font=FONT_BOLD).pack(anchor=W,padx=16,pady=(8,4))
        rf=Frame(rl,bg=C["panel"]); rf.pack(fill=X,padx=16,pady=(0,10))
        self.rep_log=Text(rf,bg=C["log_bg"],fg=C["log_txt"],font=("Consolas",9),
                          relief="flat",wrap=WORD,state=DISABLED,height=5)
        sb3=ttk.Scrollbar(rf,orient=VERTICAL,command=self.rep_log.yview)
        self.rep_log.configure(yscrollcommand=sb3.set)
        self.rep_log.pack(side=LEFT,fill=BOTH,expand=True)
        sb3.pack(side=RIGHT,fill=Y)
        self.rep_log.tag_config("ok",   foreground="#68D391")
        self.rep_log.tag_config("warn", foreground="#F6AD55")
        self.rep_log.tag_config("plain",foreground=C["log_txt"])

    # ── 공통 헬퍼 ─────────────────────────────────────────────────
    def _browse(self):
        d=filedialog.askdirectory(initialdir=self.dir_var.get())
        if d: self.dir_var.set(d)

    def _browse_rep(self):
        init=self._posts_dir or self.dir_var.get()
        d=filedialog.askdirectory(initialdir=init)
        if d: self.rep_dir_var.set(d)

    def _on_repl_scroll(self, event):
        self._repl_canvas.yview_scroll(int(-1*(event.delta/120)),"units")

    def _populate_replace_table(self, top_words):
        for w in self._repl_inner.winfo_children():
            w.destroy()
        self._replace_rows=[]
        CLF_BG={
            "의료":     ("#D9EAD3","#C6E0B8"),
            "지역·병원": ("#CFE2F3","#BDD7EE"),
            "기타":     ("#F4CCCC","#EA9999"),
        }
        for rank,(word,count) in enumerate(top_words,1):
            label=classify(word)
            bg=CLF_BG[label][rank%2]
            row=Frame(self._repl_inner,bg=bg); row.pack(fill=X)
            row.bind("<MouseWheel>",self._on_repl_scroll)
            for txt,w_ch in [(str(rank),5),(word,14),(str(count),7),(label,10)]:
                lbl=Label(row,text=txt,bg=bg,fg=C["txt"],font=FONT_SM,
                          width=w_ch,anchor="center",pady=3)
                lbl.pack(side=LEFT,padx=1)
                lbl.bind("<MouseWheel>",self._on_repl_scroll)
            sv=StringVar()
            ent=ttk.Entry(row,textvariable=sv,width=18)
            ent.pack(side=LEFT,padx=(2,4),pady=2,fill=X,expand=True)
            ent.bind("<MouseWheel>",self._on_repl_scroll)
            self._replace_rows.append((word,sv))
        self._repl_inner.update_idletasks()
        self._repl_canvas.configure(scrollregion=self._repl_canvas.bbox("all"))

    def _open_xlsx(self):
        if self._xlsx_path and os.path.exists(self._xlsx_path):
            os.startfile(self._xlsx_path)

    # ── 로그 ──────────────────────────────────────────────────────
    def _log(self, msg, tag="plain"):
        self._q.put(("log", msg, tag))

    def _append_log(self, msg, tag="plain"):
        self.log_text.configure(state=NORMAL)
        self.log_text.insert(END, msg+"\n", tag)
        self.log_text.see(END)
        self.log_text.configure(state=DISABLED)

    def _rep_log(self, msg, tag="plain"):
        def _do(m=msg, t=tag):
            self.rep_log.configure(state=NORMAL)
            self.rep_log.insert(END, m+"\n", t)
            self.rep_log.see(END)
            self.rep_log.configure(state=DISABLED)
        self.root.after(0, _do)

    # ── 진행바 ────────────────────────────────────────────────────
    def _set_step(self, step, value, maximum=100, status=""):
        self._q.put(("step", step, value, maximum, status))

    def _update_step(self, step, value, maximum, status):
        pb,stat=self.step_data[step]
        pb["maximum"]=maximum; pb["value"]=value
        if status:
            color=C["success"] if "완료" in status else C["txt2"]
            stat.configure(text=status, foreground=color)

    # ── 큐 폴링 ───────────────────────────────────────────────────
    def _poll(self):
        try:
            while True:
                item=self._q.get_nowait()
                if item[0]=="log":   self._append_log(item[1],item[2])
                elif item[0]=="step":self._update_step(item[1],item[2],item[3],item[4])
                elif item[0]=="done":self._on_done(item[1],item[2],item[3]); return
                elif item[0]=="err": self._on_error(item[1]); return
        except queue.Empty:
            pass
        if self._running: self.root.after(120,self._poll)

    # ── 추출 시작 ─────────────────────────────────────────────────
    def _start(self):
        url=self.url_var.get().strip()
        year=self.year_var.get().strip()
        out=self.dir_var.get().strip()
        if not url:
            messagebox.showwarning("입력 오류","블로그 주소를 입력해 주세요."); return
        if not year.isdigit():
            messagebox.showwarning("입력 오류","연도를 올바르게 입력해 주세요."); return
        if not out or not os.path.isdir(out):
            messagebox.showwarning("입력 오류","저장 폴더를 선택해 주세요."); return
        month_str=self.month_var.get().strip()
        month=int(month_str) if month_str.isdigit() and 1<=int(month_str)<=12 else None
        self.start_btn.configure(state=DISABLED)
        for pb,stat in self.step_data:
            pb["value"]=0; stat.configure(text="대기",foreground=C["txt2"])
        self.log_text.configure(state=NORMAL)
        self.log_text.delete("1.0",END)
        self.log_text.configure(state=DISABLED)
        self._running=True
        self.root.after(120,self._poll)
        threading.Thread(target=self._run,
                         args=(parse_blog_id(url),int(year),month,out),
                         daemon=True).start()

    # ── 백그라운드 실행 ───────────────────────────────────────────
    def _run(self, blog_id, year, month, out_dir):
        log=lambda m,t="plain": self._log(m,t)
        setp=self._set_step
        try:
            set_session_headers(blog_id)
            posts_dir=os.path.join(out_dir,"posts")
            os.makedirs(posts_dir,exist_ok=True)

            period=f"{year}년 {month}월" if month else f"{year}년 전체"
            log(f"━━━  STEP 1 · 포스트 목록 수집  ({period})  ━━━","step")
            setp(0,0,1,"수집 중…")
            posts=get_posts(blog_id,year,month,log)
            if not posts:
                self._q.put(("err","해당 연도의 포스트를 찾을 수 없습니다.")); return
            setp(0,1,1,f"완료 ({len(posts)}개)")
            log(f"  → {len(posts)}개 포스트 발견","ok")

            log("\n━━━  STEP 2 · 이미지 · 문서 다운로드  ━━━","step")
            results=[]
            for i,post in enumerate(posts,1):
                setp(1,i-1,len(posts),f"{i-1}/{len(posts)}")
                log(f"  [{i}/{len(posts)}]  {post['date']}  {post['title'][:32]}")
                results.append(process_post(blog_id,post,posts_dir,log))
                setp(1,i,len(posts),f"{i}/{len(posts)}")
                time.sleep(0.8)
            ok_cnt=sum(1 for r in results if r["status"]=="ok")
            setp(1,len(posts),len(posts),f"완료 ({ok_cnt}개)")
            log(f"  → 성공 {ok_cnt}개","ok")

            log("\n━━━  STEP 3 · 엑셀 파일 생성  ━━━","step")
            setp(2,0,3,"생성 중…")
            wb=openpyxl.Workbook(); wb.remove(wb.active)
            log("  블로그 목록 시트…")
            build_post_sheet(wb,results); setp(2,1,3)
            log("  단어 빈도 시트…")
            top=build_freq_sheet(wb,results,log); setp(2,2,3)
            month_suffix=f"_{month:02d}월" if month else ""
            xlsx_name=f"{blog_id}_{year}{month_suffix}_분석.xlsx"
            xlsx_path=os.path.join(out_dir,xlsx_name)
            wb.save(xlsx_path); setp(2,3,3,"완료")
            log(f"  → {xlsx_name}","ok")

            self._q.put(("done",xlsx_path,posts_dir,top))
        except Exception as e:
            import traceback
            self._q.put(("err",f"{e}\n\n{traceback.format_exc()}"))

    # ── 완료 처리 ─────────────────────────────────────────────────
    def _on_done(self, xlsx_path, posts_dir, top_words):
        self._running=False
        self._xlsx_path=xlsx_path
        self._posts_dir=posts_dir
        self._top_words=top_words
        self.start_btn.configure(state=NORMAL)
        self.open_xlsx_btn.configure(state=NORMAL)
        self.rep_dir_var.set(posts_dir)

        # 단어 치환 테이블 채우기
        self._populate_replace_table(top_words)

        self._log("\n✔  완료! 「결과 · 단어 치환」 탭에서 결과를 확인하세요.","ok")
        self.nb.select(1)  # 탭 2로 자동 전환
        if messagebox.askyesno("완료",
                               f"추출이 완료되었습니다.\n\n저장 위치: {os.path.dirname(xlsx_path)}\n\n"
                               "엑셀 파일을 바로 여시겠습니까?"):
            os.startfile(xlsx_path)

    def _on_error(self, msg):
        self._running=False
        self.start_btn.configure(state=NORMAL)
        self._log(f"\n✖  오류:\n{msg}","err")
        messagebox.showerror("오류",f"처리 중 오류가 발생했습니다:\n\n{msg[:300]}")

    # ── 단어 치환 ─────────────────────────────────────────────────
    def _do_replace(self):
        folder=self.rep_dir_var.get().strip()
        pairs=[(w,sv.get().strip()) for w,sv in self._replace_rows if sv.get().strip()]
        if not pairs:
            messagebox.showwarning("입력 오류","바꿀 단어를 하나 이상 입력해 주세요."); return
        if not folder or not os.path.isdir(folder):
            messagebox.showwarning("입력 오류","작업 폴더를 선택해 주세요."); return
        self.rep_log.configure(state=NORMAL)
        self.rep_log.delete("1.0",END)
        self.rep_log.configure(state=DISABLED)
        threading.Thread(target=self._run_replace,
                         args=(folder,pairs),daemon=True).start()

    def _run_replace(self, folder, pairs):
        rlog=self._rep_log
        rlog(f"{len(pairs)}개 단어 치환 시작")
        for find,repl in pairs:
            rlog(f'  "{find}"  ->  "{repl}"')
        rlog(f"\n대상 폴더: {folder}\n")
        docx_files=[]
        for root_dir,_,files in os.walk(folder):
            for f in files:
                if f.endswith(".docx"):
                    docx_files.append(os.path.join(root_dir,f))
        rlog(f"Word 파일 {len(docx_files)}개 처리 중...")
        total_replaced=0; changed_files=0
        for fp in docx_files:
            try:
                doc=Document(fp); file_cnt=0
                for para in doc.paragraphs:
                    for run in para.runs:
                        for find,repl in pairs:
                            if find in run.text:
                                run.text=run.text.replace(find,repl); file_cnt+=1
                if file_cnt>0:
                    doc.save(fp)
                    rlog(f"  v {os.path.basename(fp)}  ({file_cnt}곳)","ok")
                    changed_files+=1; total_replaced+=file_cnt
            except Exception as e:
                rlog(f"  x {os.path.basename(fp)}: {e}","warn")
        rlog(f"\n총 {changed_files}개 파일 / {total_replaced}곳 치환 완료","ok")
        if total_replaced==0:
            rlog("  -> 찾은 단어가 없습니다.","warn")

# ══════════════════════════════════════════════════════════════════
# 진입점
# ══════════════════════════════════════════════════════════════════
if __name__=="__main__":
    root=Tk()
    BlogMinerApp(root)
    root.mainloop()
