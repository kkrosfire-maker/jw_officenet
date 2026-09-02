import os
import sys
from copy import deepcopy

import docx
from docx.oxml.ns import qn
from docx.shared import Pt

CHECK, UNCHECK = "■", "□"


def resource_path(relative_path):
    base_path = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


def checklist(options, selected, sep):
    parts = []
    for opt in options:
        mark = CHECK if opt == selected else UNCHECK
        parts.append(f"{mark}{opt}")
    return sep.join(parts)


def _copy_mark_format(paragraph, run):
    # An empty template cell has no run, only a paragraph-mark <w:rPr> that
    # carries the intended font/size. A run added with python-docx inherits
    # nothing and falls back to the 11pt document default, which is why the
    # 성명 / 검사기관명 values used to print larger than their labels. Copy the
    # paragraph mark's formatting onto the new run so every field matches.
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    mark_rpr = pPr.find(qn("w:rPr"))
    if mark_rpr is None:
        return
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:rFonts", "w:sz", "w:szCs", "w:b", "w:i"):
        if rPr.find(qn(tag)) is None:
            src = mark_rpr.find(qn(tag))
            if src is not None:
                rPr.append(deepcopy(src))


def set_paragraph_text(paragraph, text):
    runs = paragraph.runs
    if not runs:
        _copy_mark_format(paragraph, paragraph.add_run(text))
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _trim_trailing_empty(cell, keep_min=1):
    """Drops the template's spare blank paragraphs from the bottom of a cell so
    the row is only as tall as its actual content (a cell must keep >= 1)."""
    while len(cell.paragraphs) > keep_min and not cell.paragraphs[-1].text.strip():
        p = cell.paragraphs[-1]._p
        p.getparent().remove(p)


def _para_max_pt(p):
    sizes = [r.font.size.pt for r in p.runs if r.font.size is not None]
    return max(sizes) if sizes else 10.0


def _tighten_paragraph(p, exact=True):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if exact:
        # An exact line height stops 한글 from applying its ~1.6x default
        # leading on import (what pushed the report onto a second page).
        # Scale it to the paragraph's own font so nothing clips.
        pf.line_spacing = Pt(_para_max_pt(p) * 1.25)
    else:
        pf.line_spacing = 1.0


def _compact_spacing(doc):
    """한글's .docx importer inflates every paragraph's leading, which is what
    pushed the report onto a second page. Tighten the spacing everywhere so the
    한글 output collapses back to one page (Word already fits)."""
    for p in doc.paragraphs:
        _tighten_paragraph(p, exact=False)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _tighten_paragraph(p, exact=True)


def _tight_cell_margins(table, tb=20, lr=80):
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblCellMar")):
        tblPr.remove(existing)
    mar = tblPr.makeelement(qn("w:tblCellMar"), {})
    for side, width in (("top", tb), ("left", lr), ("bottom", tb), ("right", lr)):
        e = mar.makeelement(qn(f"w:{side}"), {})
        e.set(qn("w:w"), str(width))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)


def _unbold_table(table):
    # The user's template carries bold sample text in a few value cells
    # (성명 / 검사자 / 판독자), which made those fields print heavier than the
    # rest. Force every run in the table back to regular weight.
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = False


def _set_cell_font_size(cell, pt):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(pt)
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None:
            mark_rpr = pPr.find(qn("w:rPr"))
            if mark_rpr is not None:
                for tag in ("w:sz", "w:szCs"):
                    e = mark_rpr.find(qn(tag))
                    if e is not None:
                        e.set(qn("w:val"), str(int(pt * 2)))


def _shrink_paragraph(p_el, pt=6):
    pPr = p_el.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = pPr.makeelement(qn("w:rPr"), {})
        pPr.append(rPr)
    for tag in ("w:sz", "w:szCs"):
        e = rPr.find(qn(tag))
        if e is None:
            e = rPr.makeelement(qn(tag), {})
            rPr.append(e)
        e.set(qn("w:val"), str(int(pt * 2)))


def _set_box_border(cell, sz=6):
    """Puts an explicit single-line border on all four sides of a cell.
    The template leaves the 기타소견 / 결론 boxes to the table style, which
    한글's .docx importer drops - so draw them on the cell itself."""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for side in ("top", "left", "bottom", "right"):
        edge = borders.makeelement(qn(f"w:{side}"), {})
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(sz))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "000000")
        borders.append(edge)
    tcPr.append(borders)


def set_cell_text(cell, text):
    set_paragraph_text(cell.paragraphs[0], text)


def set_cell_multiline(cell, text, start_index=0):
    lines = (text or "").split("\n")
    while len(cell.paragraphs) < start_index + len(lines):
        cell.add_paragraph()
    paragraphs = cell.paragraphs
    for offset, line in enumerate(lines):
        set_paragraph_text(paragraphs[start_index + offset], line)
    for p in paragraphs[start_index + len(lines):]:
        set_paragraph_text(p, "")


def detail_or_blank(detail, blank_width=18):
    if detail:
        return f"({detail})"
    return "(" + " " * blank_width + ")"


EXAM_TYPE_OPTIONS = [
    "경직장 전립선·정낭 초음파",
    "진단적 초음파",
    "진단적 초음파 (도플러 가산)",
    "제한적 초음파",
]

PRESENCE_OPTIONS = ["없음", "있음"]

TABLE3_FIELDS = [
    (1, "border", ["명확함", "불명확 또는 불규칙"]),
    (2, "symmetry", ["대칭", "비대칭"]),
    (3, "calcification", PRESENCE_OPTIONS),
    (4, "median_cyst", PRESENCE_OPTIONS),
    (5, "acute_inflammation", PRESENCE_OPTIONS),
    (6, "bladder_protrusion", PRESENCE_OPTIONS),
    (7, "bladder_stone_tumor", PRESENCE_OPTIONS),
]


def generate_docx(data, output_path, template_path=None):
    template_path = template_path or resource_path("초음파_판독지.docx")
    doc = docx.Document(template_path)

    # --- 1. 환자정보 (2 rows x 4 cols: label | value | label | value) ---
    t0 = doc.tables[0]
    set_cell_text(t0.rows[0].cells[1], data.get("reg_no", ""))
    set_cell_text(t0.rows[0].cells[3], data.get("patient_name", ""))
    set_cell_text(t0.rows[1].cells[1], data.get("birth_or_age", ""))
    set_cell_text(t0.rows[1].cells[3], checklist(["남", "여"], data.get("sex"), "  "))

    # --- 2. 검사정보 (5 rows x 4 cols: col0/col1 vertically merged) ---
    t1 = doc.tables[1]
    set_cell_text(t1.rows[0].cells[1], checklist(EXAM_TYPE_OPTIONS, data.get("exam_type"), "\n"))
    set_cell_text(t1.rows[0].cells[3], data.get("exam_date", ""))
    examiner = data.get("examiner_name", "")
    examiner_license = data.get("examiner_license", "")
    set_cell_text(
        t1.rows[1].cells[3],
        f"{examiner}    ({examiner_license})" if examiner or examiner_license else "",
    )
    set_cell_text(t1.rows[2].cells[3], data.get("read_date", ""))
    reader = data.get("reader_name", "")
    reader_license = data.get("reader_license", "")
    set_cell_text(
        t1.rows[3].cells[3],
        f"{reader}    ({reader_license})" if reader or reader_license else "",
    )
    set_cell_text(t1.rows[4].cells[3], data.get("institution", ""))

    # --- 3-(1). 필수 소견 (5 rows x 3 cols: label | value | detail) ---
    t2 = doc.tables[2]
    vol_total = data.get("vol_total", "")
    vol_transition = data.get("vol_transition", "")
    set_cell_text(t2.rows[0].cells[1], f"{vol_total} cc" if vol_total else "cc")
    set_cell_text(t2.rows[1].cells[1], f"{vol_transition} cc" if vol_transition else "cc")
    for row, field in [(2, "focal_lesion"), (3, "hyper_vascular"), (4, "seminal_vesicle")]:
        set_cell_text(t2.rows[row].cells[1], checklist(PRESENCE_OPTIONS, data.get(field), "  "))
        set_cell_text(t2.rows[row].cells[2], detail_or_blank(data.get(f"{field}_detail", "")))

    # --- 3-(2). 선택적 기술 (8 rows x 2 cols: label | checklist + detail) ---
    t3 = doc.tables[3]
    shape_options = ["삼각형", "타원형", "원형", "세로타원형"]
    set_cell_text(t3.rows[0].cells[1], checklist(shape_options, data.get("shape"), " "))
    for row, field, options in TABLE3_FIELDS:
        base = checklist(options, data.get(field), " ")
        detail = data.get(f"{field}_detail", "")
        set_cell_text(t3.rows[row].cells[1], f"{base} {detail_or_blank(detail, 8)}")

    # --- 기타소견 박스 (own 1x1 table; paragraph 0 is the fixed label) ---
    t4 = doc.tables[4]
    set_cell_multiline(t4.rows[0].cells[0], data.get("other_findings", ""), start_index=1)

    # --- 4. 결론 박스 (own 1x1 table) ---
    t5 = doc.tables[5]
    set_cell_multiline(t5.rows[0].cells[0], data.get("conclusion", ""), start_index=0)

    # Trim the template's spare blank lines and keep the free-text small so the
    # report always stays on one page - in Word and after the .hwp conversion.
    for cell in (t4.rows[0].cells[0], t5.rows[0].cells[0]):
        _trim_trailing_empty(cell, keep_min=1)
        _set_cell_font_size(cell, 8)
        _set_box_border(cell)

    trailing = doc.element.body.findall(qn("w:p"))
    if trailing:
        _shrink_paragraph(trailing[-1], pt=4)

    # Collapse the inter-line padding 한글 would otherwise add, pull the table
    # cell margins in tight (keeps the .hwp on one page too), and drop the
    # template's stray bold so every field prints at the same weight.
    _compact_spacing(doc)
    for table in doc.tables:
        _tight_cell_margins(table)
        _unbold_table(table)

    doc.save(output_path)


def convert_to_pdf(docx_path, pdf_path):
    import pythoncom
    import win32com.client as win32

    pythoncom.CoInitialize()
    word = None
    try:
        # DispatchEx always starts a brand-new Word process instead of attaching
        # to one that may already be running (e.g. a leftover instance stuck
        # behind a dialog), which is what made PDF conversion silently hang/fail
        # while the .docx (plain file write, no Word involved) always succeeded.
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        word.AutomationSecurity = 3  # msoAutomationSecurityForceDisable: suppress macro/security prompts
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        wdoc = word.Documents.Open(os.path.abspath(docx_path))
        wdoc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        wdoc.Close(False)
    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _kill_stale_hwp():
    # A 한글 instance left stuck behind its security prompt poisons every later
    # COM call ("원격 프로시저를 호출하지 못했습니다"), so clear leftovers first.
    import subprocess

    for name in ("Hwp.exe", "HwpFrame.exe"):
        try:
            subprocess.run(
                ["taskkill", "/F", "/IM", name],
                capture_output=True, check=False,
            )
        except Exception:
            pass


def convert_to_hwp(docx_path, hwp_path):
    """Opens the generated .docx in Hancom Office 한글 and re-saves it as .hwp.
    Requires 한글 (HancomOffice) to be installed on the machine."""
    import pythoncom
    import win32com.client as win32

    _kill_stale_hwp()
    pythoncom.CoInitialize()
    hwp = None
    try:
        # DispatchEx starts a fresh 한글 process (late binding, so no gen_py
        # cache write is needed inside the frozen exe).
        hwp = win32.DispatchEx("HWPFrame.HwpObject")
        try:
            # Approve automation up front so 한글 doesn't stop on its security
            # prompt when opening/saving a file it didn't create.
            hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        except Exception:
            pass
        try:
            hwp.SetMessageBoxMode(0x00020000)  # auto-dismiss modal dialogs
        except Exception:
            pass
        if os.path.exists(hwp_path):
            os.remove(hwp_path)
        # format "" lets 한글 auto-detect the .docx; forceopen skips the
        # "another program is using this file" nag.
        hwp.Open(os.path.abspath(docx_path), "", "forceopen:true")
        hwp.SaveAs(os.path.abspath(hwp_path), "HWP", "")
    finally:
        if hwp is not None:
            try:
                hwp.Clear(1)  # discard without a "save changes?" prompt
            except Exception:
                pass
            try:
                hwp.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def generate_hwp(data, output_path, template_path=None):
    """Fills the template and writes the report as a 한글 (.hwp) file.
    Returns [output_path]."""
    import tempfile

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_docx = os.path.join(tmp_dir, "report.docx")
        generate_docx(data, tmp_docx, template_path=template_path)
        convert_to_hwp(tmp_docx, output_path)
    return [output_path]


def generate_jpg(data, output_path, template_path=None, dpi=200):
    """Renders the filled report as JPG image(s). Returns the list of files written
    (more than one file when the report spans multiple pages)."""
    import tempfile

    import fitz  # PyMuPDF

    base, _ = os.path.splitext(output_path)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_docx = os.path.join(tmp_dir, "report.docx")
        tmp_pdf = os.path.join(tmp_dir, "report.pdf")
        generate_docx(data, tmp_docx, template_path=template_path)
        convert_to_pdf(tmp_docx, tmp_pdf)

        written = []
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        with fitz.open(tmp_pdf) as pdf:
            for i, page in enumerate(pdf):
                path = output_path if i == 0 else f"{base}_{i + 1}.jpg"
                pix = page.get_pixmap(matrix=matrix)
                pix.save(path)
                written.append(path)
        return written
